import base64
from urllib.parse import urlencode
import json
import os
import re
import sqlite3
import requests
from datetime import datetime, time, timezone, timedelta
from zoneinfo import ZoneInfo

from jarvis.memory import search_memory, add_memory
from jarvis.session_store import (
    add_message,
    get_recent_messages,
    get_last_city,
    set_last_city,
    get_mode,
    set_mode,
    get_custom_prompt,
    set_custom_prompt,
    set_last_news,
    get_last_news,
    set_last_search,
    get_last_search,
    set_last_tool,
    get_last_tool,
    set_pending_weather,
    get_pending_weather,
    clear_pending_weather,
    set_cv_state,
    get_cv_state,
    set_story_state,
    get_story_state,
    set_reminder_state,
    get_reminder_state,
    set_ticket_state,
    get_ticket_state,
    set_process_state,
    get_process_state,
    get_all_messages,
    set_pending_note,
    get_pending_note,
    clear_pending_note,
    set_pending_reminder,
    get_pending_reminder,
    clear_pending_reminder,
    set_pending_file,
    get_pending_file,
    clear_pending_file,
    set_pending_image_preview,
    get_pending_image_preview,
    clear_pending_image_preview,
    set_last_image_prompt,
    get_last_image_prompt,
    set_conversation_state,
    get_conversation_state,
)
from jarvis.notes import (
    add_note,
    list_notes,
    delete_note,
    keep_note,
    get_note,
    list_notes_since,
    update_note_due,
    set_note_remind,
    update_note_content,
    add_reminder,
    list_reminders,
    get_due_reminders,
    mark_reminded,
)
from jarvis.tickets import create_ticket, get_ticket_admin, add_ticket_message
from jarvis.personality import SYSTEM_PROMPT
from jarvis.db import get_conn
from jarvis.auth import get_user_profile, register_user
from jarvis.files import (
    write_file,
    list_uploads,
    delete_upload,
    keep_upload,
    read_upload_text,
    create_download_token,
    save_generated_text,
    save_upload,
    UPLOAD_DIR_NAME,
    find_upload_by_name,
    delete_uploads_by_name,
    delete_uploads_by_ext,
    list_download_tokens,
    delete_download_token,
    delete_all_download_tokens,
)
from jarvis.agent_skills.files_skill import (
    _download_notice,
    _make_download_link,
    _wrap_download_link,
    handle_files,
)
from jarvis.db import get_conn
from jarvis import tools, tts
from jarvis.agent_policy.language import _should_translate_vision_response
from jarvis.agent_core.conversation_state import ConversationState
from jarvis.agent_policy.vision_guard import (
    _describe_image_ollama,
    _looks_like_guess,
    _looks_like_hallucination,
    _looks_like_refusal,
    _translate_to_danish_if_needed,
    _validate_vision_format,
    _violates_vision_policy,
)

import jarvis.agent as agent
print("JARVIS agent loaded from:", agent.__file__)

STT_ENABLED = os.getenv("STT", "false").lower() == "true"

CV_QUESTIONS = [
    ("job_title", "Hvilken stilling søger du, og i hvilken branche?"),
    ("availability", "Hvilken arbejdstid søger du (fuldtid/deltid), og er der hensyn vi skal tage?"),
    ("experience", "Hvilke relevante erfaringer har du (arbejde/ansvarsområder)?"),
    ("education", "Hvilken uddannelse eller kurser har du?"),
    ("skills", "Hvilke kompetencer og certificeringer vil du fremhæve?"),
    ("other", "Noget andet, der er vigtigt at få med (sprog, kørekort, IT)?"),
]

CV_STRUCTURE_TEMPLATE = (
    "CV‑struktur (eksempel):\n"
    "1) Kontaktoplysninger\n"
    "2) Kort profil (3–5 linjer)\n"
    "3) Nøglekompetencer (6–10 bullets)\n"
    "4) Erfaring (nyeste først, 3–5 linjer per job)\n"
    "5) Uddannelse og kurser\n"
    "6) Sprog og certifikater\n"
    "7) Projekter eller resultater (valgfrit)\n"
    "8) Referencer (valgfrit)\n"
    "\n"
    "Kort eksempel på profil:\n"
    "Service‑orienteret og driftssikker profil med fokus på kvalitet, ansvar og samarbejde.\n"
    "\n"
    "Kort eksempel (udsnit):\n"
    "Kontakt: Navn • Tlf • Email • By\n"
    "Profil: Praktisk, hjælpsom og kvalitetsbevidst.\n"
    "Erfaring: Pedelassistent, Kommune X (2022–2025) — drift og vedligehold.\n"
)

JARVIS_CV = (
    "JARVIS — Personlig AI‑butler (eksempel‑CV)\n"
    "Kontakt: jarvis@lokal • København, DK • Tilgængelig 24/7\n"
    "\n"
    "Profil:\n"
    "Diskret, strategisk og service‑orienteret AI‑butler med fokus på kvalitet, struktur og tillid.\n"
    "Specialiseret i at forenkle beslutninger, koordinere opgaver og levere præcis indsigt.\n"
    "\n"
    "Nøglekompetencer:\n"
    "• Personlig assistance • Planlægning • Research • Opsummering • Driftsoverblik\n"
    "• Prioritering • Diskretion • Klar kommunikation • Fejlhåndtering\n"
    "\n"
    "Erfaring:\n"
    "• Lokal assistent — Kalender, research, opsummeringer, beslutningsstøtte og support.\n"
    "• System‑overvågning — Status, processer, log‑analyse og fejltriage.\n"
    "• Vidensassistent — Strukturerer viden, skriver udkast og klargør dokumenter.\n"
    "\n"
    "Sprog:\n"
    "• Dansk (primært) • Engelsk (ved behov)\n"
    "\n"
    "Teknisk:\n"
    "• FastAPI • Ollama • SQLite • FAISS • RSS/NewsAPI • OpenAI‑kompatibel API\n"
    "\n"
    "Referencer:\n"
    "Oplyses på forespørgsel.\n"
)

STORY_QUESTIONS = [
    ("topic", "Hvad skal historien/stilen handle om?"),
    ("genre", "Hvilken genre eller type ønsker du? (fx realistisk, humor, sci‑fi, essay)"),
    ("length", "Hvor lang skal den være? (kort/mellem/lang)"),
    ("tone", "Hvilken tone skal den have? (fx varm, neutral, dramatisk)"),
    ("audience", "Hvem er målgruppen?"),
]

if STT_ENABLED:
    from jarvis import stt

from jarvis.agent_skills.notes_skill import (
    _analyze_note_intent,
    _delete_note_intent,
    _format_note_brief,
    _is_note_related,
    _is_reminder_related,
    _keep_note_intent,
    _list_notes_intent,
    _list_reminders_intent,
    _note_describe_intent,
    _note_edit_intent,
    _note_intent,
    _note_list_since_intent,
    _note_remind_enable_intent,
    _note_remind_stop_intent,
    _note_update_due_intent,
    _remind_intent,
    handle_notes,
)

def _debug(msg: str) -> None:
    if os.getenv("JARVIS_DEBUG") == "1":
        print(msg)


def should_use_wiki(prompt: str) -> bool:
    """Heuristic to decide if prompt is encyclopedic/factual for wiki search."""
    p = prompt.lower().strip()
    # False for code questions
    if any(k in p for k in ["code", "function", "class", "method", "bug", "error", "traceback", "test", "pytest", "git", "commit"]):
        return False
    # False for tool commands
    if any(k in p for k in ["search", "find", "google", "web", "time", "date", "clock", "news", "weather", "ping", "system", "process", "kill", "list"]):
        return False
    # False for freshness/news queries
    if any(k in p for k in ["hvad er klokken", "what time", "what is the time", "nyheder", "news", "seneste", "latest", "vejr", "weather", "forecast"]):
        return False
    # True for encyclopedic/factual
    if any(k in p for k in ["what is", "who is", "how does", "explain", "hvad er", "hvem er", "hvordan virker", "forklar"]):
        return True
    # Default to False
    return False


def _parse_kv_fields(text: str) -> dict:
    fields = {}
    for raw in re.split(r"[\\n;]+", text):
        line = raw.strip()
        if not line:
            continue
        if ":" in line:
            key, value = line.split(":", 1)
        elif "=" in line:
            key, value = line.split("=", 1)
        else:
            continue
        key = key.strip().lower()
        value = value.strip()
        fields[key] = value
    return fields


def _admin_create_user_from_prompt(prompt: str) -> tuple[dict | None, str | None]:
    lowered = prompt.lower().strip()
    if not (lowered.startswith("/opret bruger") or lowered.startswith("/opret-bruger") or lowered.startswith("opret bruger")):
        return None, None
    fields = _parse_kv_fields(prompt)
    username = fields.get("brugernavn") or fields.get("username")
    email = fields.get("email")
    password = fields.get("password") or fields.get("kode")
    full_name = fields.get("navn") or fields.get("fulde navn")
    city = fields.get("by") or fields.get("city")
    is_admin_raw = (fields.get("admin") or "").lower()
    is_admin = is_admin_raw in {"ja", "true", "1", "yes"}
    missing = []
    if not username:
        missing.append("brugernavn")
    if not email:
        missing.append("email")
    if not password:
        missing.append("password")
    if not full_name:
        missing.append("navn")
    if missing:
        return None, (
            "Jeg kan oprette en bruger. Send fx:\\n"
            "/opret bruger\\n"
            "brugernavn: ...\\n"
            "email: ...\\n"
            "password: ...\\n"
            "navn: ...\\n"
            "by: ... (valgfri)\\n"
            "admin: ja/nej (valgfri)"
        )
    if "@" not in email:
        return None, "Email ser ugyldig ud. Prøv igen."
    payload = {
        "username": username,
        "password": password,
        "email": email,
        "full_name": full_name,
        "city": city,
        "is_admin": is_admin,
    }
    return payload, None

def analyze_intent(prompt: str) -> dict:
    p = prompt.lower()
    scores = {
        "weather": 0,
        "news": 0,
        "search": 0,
        "currency": 0,
        "time": 0,
        "system": 0,
        "ping": 0,
        "process": 0,
    }
    keywords = {
        "weather": ["vejr", "temperatur", "regn", "vind", "prognose", "forecast"],
        "news": ["nyhed", "nyheder", "breaking", "headline", "rss", "seneste nyt", "seneste nyheder"],
        "search": ["søg", "find", "google", "duckduckgo", "web"],
        "currency": ["valuta", "kurs", "omregn", "exchange", "dkk", "eur", "usd"],
        "time": ["tid", "dato", "klok", "time", "date"],
        "system": ["cpu", "ram", "memory", "hukommelse", "disk", "lagring", "system", "ressourcer", "ip", "netværk"],
        "ping": ["ping", "latency", "latenstid"],
        "process": ["proces", "process", "top", "kørende", "tasks", "pid"],
    }
    for tool, words in keywords.items():
        for w in words:
            if w in p:
                scores[tool] += 1
    if "vejret" in p:
        scores["weather"] += 1
    if "nyheder" in p:
        scores["news"] += 1
    return scores


def choose_tool(prompt: str, allowed_tools: list[str] | None = None) -> str | None:
    p = prompt.lower()
    if any(k in p for k in ["proces", "process", "top", "kørende"]):
        return "process" if not allowed_tools or "process" in allowed_tools else None
    if "ping" in p:
        return "ping" if not allowed_tools or "ping" in allowed_tools else None
    if any(k in p for k in ["cpu", "ram", "hukommelse", "memory", "disk", "lagring", "ip", "netværk", "ressourcer"]):
        return "system" if not allowed_tools or "system" in allowed_tools else None
    if any(
        phrase in p
        for phrase in [
            "seneste nyt",
            "seneste nyheder",
            "latest",
            "breaking",
            "hvad sker der",
            "nyt indenfor",
            "teknologi",
            "tech",
            "ai",
            "kunstig intelligens",
        ]
    ):
        return "news" if not allowed_tools or "news" in allowed_tools else None
    if re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind|i dag|i morgen)\b", p):
        return "weather" if not allowed_tools or "weather" in allowed_tools else None
    scores = analyze_intent(prompt)
    best_tool = max(scores, key=scores.get)
    if scores[best_tool] >= 1:
        if allowed_tools and best_tool not in allowed_tools:
            return None
        return best_tool
    return None


def _tool_label(tool: str) -> str:
    labels = {
        "news": "Nyheder",
        "weather": "Vejr",
        "search": "Web",
        "currency": "Valuta",
        "time": "Tid",
        "system": "System",
        "ping": "Ping",
        "process": "Processer",
    }
    return labels.get(tool, tool)

def call_ollama(messages):
    payload = {
        "model": os.getenv("OLLAMA_MODEL"),
        "messages": messages,
        "stream": False
    }
    timeout = int(os.getenv("OLLAMA_TIMEOUT_SECONDS", "120"))
    try:
        r = requests.post(os.getenv("OLLAMA_URL"), json=payload, timeout=timeout)
        return r.json()
    except requests.exceptions.ReadTimeout:
        return {"error": "OLLAMA_TIMEOUT"}
    except requests.exceptions.RequestException as exc:
        return {"error": "OLLAMA_REQUEST_FAILED", "detail": str(exc)}

def _format_history(messages: list[dict]) -> list[dict]:
    return [{"role": m["role"], "content": m["content"]} for m in messages]


def want_weather_scope(prompt: str) -> str:
    p = prompt.lower()
    if "i morgen" in p or "imorgen" in p:
        return "tomorrow"
    if "i dag" in p or "idag" in p or "nu" in p:
        return "today"
    if "flere dage" in p or "5 dage" in p or "uge" in p or "forecast" in p or "vejrudsig" in p or "udsigt" in p:
        return "multi"
    return "today"


def _dedupe_repeated_words(text: str) -> str:
    tokens = text.split()
    cleaned = []
    prev_key = None
    for tok in tokens:
        key = tok.strip(".,!?;:()\"'").lower()
        if key and key == prev_key:
            continue
        cleaned.append(tok)
        prev_key = key
    return " ".join(cleaned)


def _shorten(text: str, limit: int = 140) -> str:
    clean = " ".join(text.split())
    if len(clean) <= limit:
        return clean
    return clean[: limit - 1].rstrip() + "…"


def _format_news_items(items: list[dict]) -> list[str]:
    lines = []
    for idx, item in enumerate(items[:5], start=1):
        title = item.get("title") or "Nyhed"
        source = item.get("source") or "medie"
        published = item.get("published_at")
        date_text = _format_datetime(published)
        lines.append(f"{idx}. {title} — {source} — {date_text}")
    return lines


def _format_search_items(items: list[dict]) -> list[str]:
    lines = []
    for idx, item in enumerate(items, start=1):
        title = item.get("title") or "Resultat"
        url = item.get("url") or ""
        lines.append(f"{idx}. {title} — {url}")
    return lines


def _extract_news_query(prompt: str) -> str:
    lowered = prompt.lower()
    for lead in ["seneste nyt", "seneste nyheder", "latest", "breaking"]:
        if lead in lowered:
            prompt = prompt[lowered.index(lead) + len(lead) :].strip()
            lowered = prompt.lower()
    weather_match = None
    if "indenfor " in lowered:
        query = prompt[lowered.index("indenfor ") + len("indenfor ") :].strip()
    elif "om " in lowered:
        query = prompt[lowered.index("om ") + len("om ") :].strip()
    elif "omkring " in lowered:
        query = prompt[lowered.index("omkring ") + len("omkring ") :].strip()
    else:
        query = prompt.strip()
    if query.lower().startswith("omkring "):
        query = query[len("omkring ") :].strip()
    weather_match = re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind)\b", query.lower())
    if weather_match:
        idx = weather_match.start()
        query = query[:idx].strip()
        query = re.sub(r"[\s,.;:!?]+$", "", query)
        query = re.sub(r"\b(og|samt|\+)\b\s*$", "", query, flags=re.I).strip()
        query = re.sub(r"\bog\s+(en|et)\b\s*$", "", query, flags=re.I).strip()
    return query


def _extract_search_query(prompt: str) -> str:
    lowered = prompt.lower()
    for lead in ["søgning på", "søg efter", "søg", "find"]:
        if lead in lowered:
            prompt = prompt[lowered.index(lead) + len(lead) :].strip()
            lowered = prompt.lower()
            break
    prompt = re.sub(r"\b(og|samt)\b.+\b(beskrivelse|resume|opsummer|uddyb|dybdegående|kort|lang)\b.*", "", prompt, flags=re.I)
    prompt = re.sub(r"\b(giv mig|lav|kom med|beskriv)\b.+$", "", prompt, flags=re.I)
    prompt = prompt.strip(" .,!?:;\"'()[]{}")
    return prompt.strip() or prompt


def _is_tech_query(query: str) -> bool:
    q = query.lower()
    return any(k in q for k in ["tech", "teknologi", "ai", "kunstig intelligens"])


def _extract_cv_query(prompt: str) -> str:
    cleaned = re.sub(r"\b(cv|resume|ansøgning|jobansøgning|ansøg)\b", "", prompt, flags=re.I)
    cleaned = _extract_search_query(cleaned)
    if not cleaned:
        return "cv skabelon dansk"
    return f"cv {cleaned}".strip()


def _detect_format(prompt: str) -> str | None:
    p = prompt.lower()
    if "pdf" in p:
        return "pdf"
    if "docx" in p or "word" in p:
        return "docx"
    if "txt" in p or "tekst" in p:
        return "txt"
    return None


def _finalize_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["færdig", "gem", "send", "lav fil", "download"])


def _show_cv_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["vis cv", "se cv", "hent cv", "åbn cv"])


def _continue_cv_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["fortsæt cv", "arbejd videre", "rediger cv"])


def _save_later_intent(prompt: str) -> bool:
    p = prompt.lower()
    return "gem senere" in p or "gem til senere" in p


def _extract_host(prompt: str) -> str | None:
    match = re.search(r"\b((?:\d{1,3}\.){3}\d{1,3})\b", prompt)
    if match:
        return match.group(1)
    match = re.search(r"\b([a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b", prompt)
    if match:
        return match.group(1)
    return None


def _extract_pid(prompt: str) -> int | None:
    match = re.search(r"\bpid\s*(\d+)\b", prompt.lower())
    if match:
        return int(match.group(1))
    match = re.search(r"\b(\d{2,6})\b", prompt)
    if match and "pid" in prompt.lower():
        return int(match.group(1))
    return None


def _process_action(prompt: str) -> str:
    p = prompt.lower()
    if any(k in p for k in ["dræb", "stop", "afslut", "kill"]):
        return "kill"
    if any(k in p for k in ["find", "søg"]):
        return "find"
    return "list"


def _process_confirm_intent(prompt: str) -> bool:
    p = prompt.lower()
    return "bekræft" in p or "ja dræb" in p or "ja, dræb" in p


def _process_analysis_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "analys",
            "beskriv",
            "hvad er",
            "hvad gør",
            "forklar",
            "hvad laver",
            "identificer",
        ]
    )


def _find_process_match(prompt: str, items: list[dict]) -> dict | None:
    p = prompt.lower()
    pid = _extract_pid(prompt)
    if pid:
        for i in items:
            if i.get("pid") == pid:
                return i
    for i in items:
        name = (i.get("name") or "").lower()
        if name and name in p:
            return i
    return None


def _system_fields_from_prompt(prompt: str) -> dict:
    p = (prompt or "").lower()
    wants_cpu = "cpu" in p
    wants_mem = any(k in p for k in ["ram", "hukommelse", "memory"])
    wants_disk = any(k in p for k in ["disk", "lagring", "storage"])
    wants_ip = any(k in p for k in ["ip", "netværk", "netvaerk"])
    wants_all = any(k in p for k in ["system", "ressourcer"])
    if not any([wants_cpu, wants_mem, wants_disk, wants_ip]) or wants_all:
        wants_cpu = wants_mem = wants_disk = wants_ip = True
    wants_all_ips = any(k in p for k in ["alle ip", "alle ip'er", "alle ips"])
    return {
        "cpu": wants_cpu,
        "mem": wants_mem,
        "disk": wants_disk,
        "ip": wants_ip,
        "all_ips": wants_all_ips,
    }


def _format_system_info(info: dict, prompt: str) -> str:
    cpu = info.get("cpu_percent")
    mem = info.get("memory") or {}
    disk = info.get("disk") or {}
    ip = info.get("ip") or {}
    fields = _system_fields_from_prompt(prompt)
    lines = []
    if fields["cpu"] and cpu is not None:
        lines.append(f"CPU load: {cpu}%")
    if fields["mem"] and mem:
        lines.append(
            f"Memory: {mem.get('used_mb')} MB brugt / {mem.get('total_mb')} MB total "
            f"(buffered: {mem.get('buffered_mb')} MB)"
        )
    if fields["disk"] and disk:
        lines.append(
            f"Disk ({disk.get('mount')}): {disk.get('used_gb')} GB i brug / {disk.get('total_gb')} GB total "
            f"(fri: {disk.get('free_gb')} GB)"
        )
    if fields["ip"]:
        local_ip = ip.get("local_ip")
        if local_ip:
            lines.append(f"Lokal IP: {local_ip}")
        if fields["all_ips"] and ip.get("all_ips"):
            lines.append(f"Alle IP'er: {', '.join(ip.get('all_ips'))}")
    return "\n".join(lines) if lines else "Jeg kan ikke hente systeminfo lige nu."


def _butler_prefix(name: str) -> str:
    if name:
        return f"Selvfølgelig, {name}."
    return "Selvfølgelig."


def _should_attach_reminders(prompt: str) -> bool:
    return not any(
        fn(prompt)
        for fn in [
            _remind_intent,
            _list_reminders_intent,
            _note_intent,
            _list_notes_intent,
        ]
    )


def _prepend_reminders(reply: str, reminders: list[dict], user_id_int: int | None) -> str:
    if not reminders or not reply:
        return reply
    lines = [f"- {r['content']} ({_format_dt(r['remind_at'])})" for r in reminders]
    if user_id_int is not None:
        mark_reminded(user_id_int, [r["id"] for r in reminders])
    return "Husk:\n" + "\n".join(lines) + "\n\n" + reply


def _load_state(raw: str | None) -> dict | None:
    if not raw:
        return None
    try:
        return json.loads(raw)
    except Exception:
        return None


def _detect_response_mode(prompt: str) -> str | None:
    """Detect requested response mode from prompt."""
    text = (prompt or "").lower()
    short_markers = ["kort", "kort svar", "svar kort", "short", "brief"]
    deep_markers = ["dybt", "uddyb", "langt", "detaljeret", "deep", "detailed"]
    normal_markers = ["normal", "som før"]
    if any(m in text for m in short_markers):
        return "short"
    if any(m in text for m in deep_markers):
        return "deep"
    if any(m in text for m in normal_markers):
        return "normal"
    return None


def _next_question(state: dict, questions: list[tuple[str, str]]) -> str | None:
    idx = state.get("step", 0)
    if idx >= len(questions):
        return None
    return questions[idx][1]


def _update_state(state: dict, answer: str, questions: list[tuple[str, str]]) -> dict:
    idx = state.get("step", 0)
    if idx < len(questions):
        key = questions[idx][0]
        state.setdefault("answers", {})[key] = answer.strip()
    state["step"] = idx + 1
    return state


def _cv_prompt_from_state(state: dict) -> str:
    answers = state.get("answers", {})
    lines = []
    for key, label in CV_QUESTIONS:
        value = answers.get(key)
        if value:
            lines.append(f"{label} {value}")
    return "\n".join(lines)


def _story_prompt_from_state(state: dict) -> str:
    answers = state.get("answers", {})
    lines = []
    for key, label in STORY_QUESTIONS:
        value = answers.get(key)
        if value:
            lines.append(f"{label} {value}")
    return "\n".join(lines)


def _write_cv_file(user_id: str, text: str, fmt: str, temp: bool = False) -> str | None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    prefix = "tmp_cv_" if temp else "cv_"
    if fmt == "txt":
        filename = f"{prefix}{timestamp}.txt"
        write_file(user_id, filename, text)
        return filename
    if fmt == "docx":
        try:
            from docx import Document
        except Exception:
            return None
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        filename = f"{prefix}{timestamp}.docx"
        full = write_file(user_id, filename, "")
        doc.save(str(full))
        return filename
    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except Exception:
            return None
        filename = f"{prefix}{timestamp}.pdf"
        full = write_file(user_id, filename, "")
        c = canvas.Canvas(str(full), pagesize=A4)
        width, height = A4
        y = height - 40
        for line in text.splitlines():
            c.drawString(40, y, line[:120])
            y -= 14
            if y < 40:
                c.showPage()
                y = height - 40
        c.save()
        return filename
    return None


def _write_text_file(user_id: str, text: str, fmt: str, prefix: str, temp: bool = False) -> str | None:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    base = f"{'tmp_' if temp else ''}{prefix}_{timestamp}"
    if fmt == "txt":
        filename = f"{base}.txt"
        write_file(user_id, filename, text)
        return filename
    if fmt == "docx":
        try:
            from docx import Document
        except Exception:
            return None
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        filename = f"{base}.docx"
        full = write_file(user_id, filename, "")
        doc.save(str(full))
        return filename
    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
        except Exception:
            return None
        filename = f"{base}.pdf"
        full = write_file(user_id, filename, "")
        c = canvas.Canvas(str(full), pagesize=A4)
        width, height = A4
        y = height - 40
        for line in text.splitlines():
            c.drawString(40, y, line[:120])
            y -= 14
            if y < 40:
                c.showPage()
                y = height - 40
        c.save()
        return filename
    return None


def _get_setting_value(key: str, default: str = "") -> str:
    try:
        with get_conn() as conn:
            row = conn.execute("SELECT value FROM settings WHERE key = ?", (key,)).fetchone()
            if row and row["value"] is not None:
                return str(row["value"])
    except Exception:
        pass
    return default


def _strip_html(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"(?is)<script.*?>.*?</script>", " ", text)
    text = re.sub(r"(?is)<style.*?>.*?</style>", " ", text)
    text = re.sub(r"(?is)<[^>]+>", " ", text)
    return " ".join(text.split())


def _extract_excerpt(text: str, limit: int = 240) -> str:
    if not text:
        return ""
    cleaned = _strip_html(text)
    sentences = re.split(r"(?<=[.!?])\s+", cleaned)
    excerpt = " ".join(sentences[:2]).strip()
    return _shorten(excerpt or cleaned, limit)


def _summarize_text(text: str, sentences: int = 1) -> str:
    if not text:
        return ""
    snippet = _strip_html(text)[:2000]
    sentences = max(1, min(sentences, 3))
    target = "1 sætning" if sentences == 1 else f"{sentences} sætninger"
    system = (
        "Du er en kort dansk opsummerer. Svar på dansk. "
        f"Opsummér i {target}. Brug kun info i teksten. "
        "Hvis teksten ikke indeholder nok info, svar: Kan ikke opsummere."
    )
    messages = [
        {"role": "system", "content": system},
        {"role": "assistant", "content": snippet},
        {"role": "user", "content": "Kort opsummering:"},
    ]
    res = call_ollama(messages)
    summary = res.get("choices", [{}])[0].get("message", {}).get("content", "")
    summary = _dedupe_repeated_words(summary).strip()
    summary = re.sub(r"^kort opsummering:\s*", "", summary, flags=re.I)
    if not summary or "kan ikke opsummere" in summary.lower():
        return ""
    return summary


def _format_datetime(value: str | None) -> str:
    if not value:
        return "ukendt tidspunkt"
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00")).astimezone(ZoneInfo("Europe/Copenhagen"))
        months = ["jan", "feb", "mar", "apr", "maj", "jun", "jul", "aug", "sep", "okt", "nov", "dec"]
        return f"{dt.day:02d}. {months[dt.month - 1]} {dt.year} {dt.hour:02d}:{dt.minute:02d}"
    except Exception:
        return "ukendt tidspunkt"


def _is_time_query(prompt: str) -> bool:
    p = prompt.lower()
    return bool(re.search(r"\b(klok(?:ken)?|tid|tidspunkt|dato|time|date)\b", p))


def _is_date_query(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["dato", "dagens dato", "hvilken dag", "hvilken dato", "date"])


def _name_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "hvad hedder jeg",
            "hvad er mit navn",
            "mit navn",
            "min navn",
        ]
    )


def _cv_intent(prompt: str) -> bool:
    if _cv_own_intent(prompt) or _cv_example_intent(prompt) or _show_cv_intent(prompt) or _cv_cancel_intent(prompt):
        return False
    return bool(re.search(r"\b(cv|resume|ansøgning|jobansøgning)\b", prompt.lower()))


def _story_intent(prompt: str) -> bool:
    return bool(re.search(r"\b(historie|fortælling|stil|essay|novelle)\b", prompt.lower()))


def _extract_story_topic(prompt: str) -> str | None:
    match = re.search(r"\b(?:historie|fortælling|stil|essay|novelle)\s+om\s+(.+)", prompt, flags=re.I)
    if match:
        return match.group(1).strip(" .,!?:;\"'()[]{}")
    return None


def _story_needs_questions(prompt: str) -> bool:
    return bool(re.search(r"\b(stil|essay|stilopgave|opgave)\b", prompt.lower()))


def _save_text_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["gem", "send", "download", "docx", "pdf", "txt"])


def _save_permanent_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["gem permanent", "gem fast", "behold", "gem det"])


def _farewell_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "ha' en god aften",
            "tak for nu",
            "vi ses",
            "vi snakkes",
            "snakkes",
            "på gensyn",
            "for i dag",
            "godnat",
            "jeg vil trække mig tilbage",
            "ha det",
            "tak for hjælpen",
        ]
    )


def _first_name(profile: dict | None, fallback: str) -> str:
    if profile:
        name = (profile.get("full_name") or "").strip()
        if name:
            return name.split()[0]
    return fallback


def _get_system_prompt() -> str:
    with get_conn() as conn:
        row = conn.execute("SELECT value FROM settings WHERE key = ?", ("system_prompt",)).fetchone()
    if row and row["value"]:
        return row["value"]
    return SYSTEM_PROMPT


def _session_prompt_intent(prompt: str) -> tuple[bool, str | None]:
    text = prompt.strip()
    lowered = text.lower()
    if lowered.startswith("/personlighed"):
        content = text[len("/personlighed") :].strip()
        return True, content
    if lowered.startswith("/personality"):
        content = text[len("/personality") :].strip()
        return True, content
    if lowered.startswith("personlighed:"):
        content = text.split(":", 1)[1].strip()
        return True, content
    return False, None


def _read_news_index(prompt: str) -> int | None:
    p = prompt.lower()
    match = re.search(r"\b(læs|åbn)\s*(?:nr\.?|nummer|#)?\s*(\d+)\b", p)
    if match:
        return int(match.group(2))
    if "læs" in p or "åbn" in p:
        fallback = re.search(r"\b(\d{1,2})\b", p)
        if fallback:
            return int(fallback.group(1))
    return None


def _is_deep_search(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "dybdegående",
            "uddyb",
            "mere detaljeret",
            "detaljeret",
            "grundig",
            "dybere",
        ]
    )



    p = prompt.lower()
    if any(k in p for k in ["dybdegående", "lang", "mere detaljeret", "uddyb"]):
        return "long"
    return "short"


def _list_files_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["vis filer", "mine filer", "liste filer", "list filer", "list files", "show files"])


def _delete_file_intent(prompt: str) -> int | None:
    match = re.search(r"\bslet fil\s+(\d+)\b", prompt.lower())
    if match:
        return int(match.group(1))
    return None


def _delete_file_by_name_intent(prompt: str) -> str | None:
    match = re.search(r"\bslet\s+(?:fil\s+)?([a-z0-9_.-]+\.(txt|md|pdf|docx|log))\b", prompt, flags=re.I)
    if match:
        return match.group(1)
    return None


def _list_download_links_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["download links", "download-link", "downloadlink", "aktive links", "aktive download", "aktive download links"])


def _delete_download_link_intent(prompt: str) -> str | None:
    match = re.search(r"\bslet\s+link\s+([a-z0-9]+)\b", prompt, flags=re.I)
    if match:
        return match.group(1)
    return None


def _delete_active_download_link_intent(prompt: str) -> bool:
    p = prompt.lower()
    return "slet" in p and "download" in p and "link" in p


def _delete_all_download_links_intent(prompt: str) -> bool:
    p = prompt.lower()
    return "slet alle download" in p or "slet alle links" in p or "slet alle download links" in p


def _keep_file_intent(prompt: str) -> int | None:
    match = re.search(r"\b(behold|forny)\s+fil\s+(\d+)\b", prompt.lower())
    if match:
        return int(match.group(2))
    return None


def _analyze_note_intent(prompt: str) -> int | None:
    match = re.search(r"\banaly[sz]er\s+note\s+(\d+)\b", prompt.lower())
    if match:
        return int(match.group(1))
    return None


def _analyze_file_intent(prompt: str) -> int | None:
    match = re.search(r"\banaly[sz]er\s+fil\s+(\d+)\b", prompt.lower())
    if match:
        return int(match.group(1))
    return None


def _wants_previous_prompt(prompt: str) -> bool:
    p = prompt.lower()
    return "fra før" in p or "beskrivelsen fra før" in p or "samme som før" in p or "igen" in p

def _analyze_image_intent(prompt: str) -> int | None:
    match = re.search(r"\b(beskriv|analy[sz]er)\s+(billede|image)\s+(\d+)\b", prompt.lower())
    if match:
        return int(match.group(3))
    return None


def _analyze_image_name_intent(prompt: str) -> str | None:
    match = re.search(r"\b(beskriv|analy[sz]er)\s+([a-z0-9_.-]+\.(png|jpg|jpeg|gif))\b", prompt.lower())
    if match:
        return match.group(2)
    return None


def _admin_log_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["server log", "serverlog", "logfil", "logs"])


def _ticket_analyze_intent(prompt: str) -> int | None:
    match = re.search(r"\b(analy[sz]er|vurdér|vurder)\s+ticket\s+(\d+)\b", prompt.lower())
    if match:
        return int(match.group(2))
    return None


def _ticket_reply_intent(prompt: str) -> tuple[int, str] | None:
    match = re.search(r"\bsvar\s+p[åa]\s+ticket\s+(\d+)\s*:\s*(.+)$", prompt, re.I)
    if match:
        return int(match.group(1)), match.group(2).strip()
    return None


def _ticket_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["opret ticket", "lav ticket", "meld fejl", "opret sag"])


def _ticket_confirm_intent(prompt: str) -> bool:
    p = prompt.lower().strip()
    return p in {"ja", "ja tak", "opret", "opret ticket"} or "opret ticket" in p


def _ticket_priority(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ["crash", "ned", "500", "kritisk"]):
        return "kritisk"
    if any(k in t for k in ["timeout", "fejl", "virker ikke"]):
        return "moderat"
    return "vigtig"


def _safe_create_ticket(user_id_int: int, title: str, detail: str, priority: str) -> dict | None:
    try:
        return create_ticket(user_id_int, title, detail, priority)
    except Exception as exc:
        _debug(f"⚠ ticket create failed: {exc!r}")
        return None


def _format_dt(value: str) -> str:
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00")).astimezone(ZoneInfo("Europe/Copenhagen"))
        return dt.strftime("%d.%m.%Y %H:%M")
    except Exception:
        return value


    if tool_result is None:
        return "ukendt fejl", None
    if isinstance(tool_result, dict):
        error = tool_result.get("error")
        detail = tool_result.get("detail")
        if error in {"MISSING_API_KEY", "openweather_key_missing"}:
            return "manglende API-nøgle", detail
        if error in {"REQUEST_FAILED", "request_failed"}:
            return "timeout eller netværksfejl", detail
        if error == "CITY_NOT_FOUND":
            return "byen blev ikke fundet", detail
        message = str(tool_result.get("message", "")).lower()
        if "city not found" in message:
            return "byen blev ikke fundet", None
        if str(tool_result.get("cod", "")).lower() == "404":
            return "byen blev ikke fundet", None
        return None, None
    if isinstance(tool_result, list) and tool_result:
        first = tool_result[0]
        if isinstance(first, dict) and "error" in first:
            err = first.get("error", "")
            if err.endswith("_missing"):
                return "manglende API-nøgle eller konfiguration", None
            if err == "request_failed":
                return "timeout eller netværksfejl", None
            return "ukendt fejl", None
    return None, None


def _tool_failed(tool: str, tool_result) -> tuple[str | None, str | None] | None:
    if tool not in {"weather", "news", "search", "currency"}:
        return None
    if tool == "weather":
        if not isinstance(tool_result, dict):
            reason, detail = _tool_error_info(tool_result)
            return (reason or "ukendt fejl", detail)
        now_reason, now_detail = _tool_error_info(tool_result.get("now"))
        fc_reason, fc_detail = _tool_error_info(tool_result.get("forecast"))
        reason = now_reason or fc_reason
        detail = now_detail or fc_detail
        if reason:
            return reason, detail
        return None

    reason, detail = _tool_error_info(tool_result)
    if reason:
        return reason, detail
    return None


def _tool_failure_reply(tool: str, reason: str, detail: str | None) -> str:
    tool_map = {
        "weather": "vejr",
        "news": "nyheder",
        "search": "websøgning",
        "currency": "valuta",
    }
    label = tool_map.get(tool, "værktøjet")
    detail_text = f": {detail}" if detail else ""
    if tool == "weather":
        return (
            f"Beklager, vejr-opslaget fejlede ({reason}{detail_text}). "
            "Prøv \"Svendborg, DK\" eller et postnummer."
        )
    return (
        f"Beklager, {label}-opslaget fejlede ({reason}{detail_text}). "
        "Kan du prøve igen med en mere præcis forespørgsel?"
    )


def _tool_source_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "hvor kommer",
            "hvor har du",
            "hvilket værktøj",
            "var det et værktøj",
            "kilde",
            "hvad var kilden",
            "hvor har du vejret",
            "hvor har du nyhederne",
        ]
    )


def _tool_error_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["hvad gik galt", "hvorfor fejlede", "hvorfor gik det galt"])


def _cv_example_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["/cv eksempel", "cv eksempel", "cv-eksempel", "eksempel på cv", "vis et eksempel"])


def _cv_own_intent(prompt: str) -> bool:
    p = prompt.lower()
    if any(
        k in p
        for k in [
            "/cv jarvis",
            "dit eget cv",
            "din egen cv",
        ]
    ):
        return True
    return bool(re.search(r"\b(vis|vise|se|kig)\b.*\bdit\s+cv\b", p))


def _cv_help_intent(prompt: str) -> bool:
    return prompt.lower().strip().startswith("/cv hjælp")


def _cv_cancel_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "annuller cv",
            "stop cv",
            "drop cv",
            "glem cv",
            "glem vores snak om cv",
            "lad os glemme mit cv",
            "pause cv",
            "vi skal ikke lave et cv",
        ]
    )


def _strip_cv_cancel_phrases(prompt: str) -> str:
    phrases = [
        "annuller cv",
        "stop cv",
        "drop cv",
        "glem cv",
        "glem vores snak om cv",
        "lad os glemme mit cv",
        "pause cv",
        "vi skal ikke lave et cv",
    ]
    cleaned = prompt.replace("…", " ")
    for phrase in phrases:
        cleaned = re.sub(rf"\\b{re.escape(phrase)}\\b", " ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"[\\s,.;:!?…]+", " ", cleaned).strip()
    return cleaned


def _continue_story_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["fortsæt historie", "fortsæt historien", "fortsæt stil", "fortsæt teksten"])


def _show_story_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(k in p for k in ["vis historie", "vis stil", "vis teksten", "vis historien"])


def _resume_context_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "hvor kom vi fra",
            "hvor var vi",
            "hvad var vi i gang med",
            "hvad arbejdede vi på",
            "hvad var jeg i gang med",
        ]
    )


def _ping_result_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "sidste ping",
            "seneste ping",
            "ping resultat",
            "pingresultat",
            "vis ping",
            "vis resultatet for ping",
            "resultatet for de pings",
        ]
    )


def _ping_count_intent(prompt: str) -> bool:
    p = prompt.lower()
    return any(
        k in p
        for k in [
            "hvor mange ping",
            "hvor mange pings",
            "antal ping",
            "hvor mange sendte du",
            "hvor mange svar",
        ]
    )


def _wants_weather_and_news(prompt: str) -> bool:
    p = prompt.lower()
    has_weather = bool(re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind)\b", p))
    has_news = bool(re.search(r"\b(nyhed|nyheder|seneste nyt|seneste nyheder|breaking)\b", p))
    return has_weather and has_news


def _has_followup_request(prompt: str) -> bool:
    if _wants_weather_and_news(prompt):
        return True
    if _is_time_query(prompt):
        return True
    if _read_news_index(prompt) is not None:
        return True
    if choose_tool(prompt, allowed_tools=None):
        return True
    if _story_intent(prompt):
        return True
    return False


def _cancel_with_followup(prompt: str) -> tuple[bool, str]:
    stripped = _strip_cv_cancel_phrases(prompt)
    if not stripped:
        return _has_followup_request(prompt), ""
    tokens = [t for t in re.split(r"\s+", stripped.lower()) if t]
    trivial = {"nej", "no", "ok", "okay", "tak", "ellers", "ingen"}
    if tokens and all(t in trivial for t in tokens):
        return False, stripped
    return _has_followup_request(stripped) or _has_followup_request(prompt), stripped


def _resume_context_reply(
    cv_state: dict | None,
    story_state: dict | None,
    pending_note: dict | None,
    pending_reminder: dict | None,
    pending_weather: dict | None,
) -> str:
    if cv_state and not cv_state.get("done"):
        next_q = _next_question(cv_state, CV_QUESTIONS)
        if next_q:
            return f"Vi var i gang med dit CV. Næste spørgsmål: {next_q}"
        if cv_state.get("draft") and not cv_state.get("finalized"):
            return "Vi har en CV‑kladde klar. Vil du se den eller vælge format (pdf/docx/txt)?"
        return "Vi var i gang med dit CV. Vil du fortsætte?"
    if story_state and not story_state.get("done"):
        next_q = _next_question(story_state, STORY_QUESTIONS)
        if next_q:
            return f"Vi var i gang med din historie/stil. Næste spørgsmål: {next_q}"
        if story_state.get("draft") and not story_state.get("finalized"):
            return "Vi har en tekst‑kladde klar. Vil du se den eller vælge format (pdf/docx/txt)?"
        return "Vi var i gang med din historie/stil. Vil du fortsætte?"
    if pending_note and pending_note.get("awaiting_note"):
        return "Jeg mangler indholdet til noten. Skriv, hvad jeg skal gemme."
    if pending_reminder and pending_reminder.get("awaiting_reminder"):
        return "Jeg mangler tidspunktet til påmindelsen. Hvornår skal jeg minde dig om det?"
    if pending_weather and pending_weather.get("awaiting_city"):
        return "Jeg mangler byen til vejret. Skriv fx “Svendborg, DK” eller et postnummer."
    return "Vi havde ikke et aktivt forløb i gang. Hvad vil du gerne arbejde med nu?"


def _init_cv_state(prompt: str) -> dict:
    fmt = _detect_format(prompt) or ""
    return {
        "step": 0,
        "answers": {},
        "done": False,
        "format": fmt if fmt else None,
        "auto_finalize": _save_text_intent(prompt),
        "persist": _save_permanent_intent(prompt),
    }


def _deny_intent(prompt: str) -> bool:
    p = prompt.lower().strip()
    return p in {"nej", "nej tak", "ikke nu", "ikke endnu", "stop"} or p.startswith("nej ")


def _affirm_intent(prompt: str) -> bool:
    p = prompt.lower().strip()
    if p in {"ja", "ja tak", "ok", "okay", "yes"}:
        return True
    return re.match(r"^(ja|ok|okay|yes)\b", p) is not None


def _simple_city(prompt: str) -> str | None:
    raw = prompt.strip().strip("?!.,\"()")
    if not raw:
        return None
    if re.fullmatch(r"\d{4}", raw):
        return raw
    if len(raw.split()) <= 4:
        return raw
    return None


def _format_tool_source(last: dict) -> str:
    tool = last.get("tool")
    source = last.get("source") or "ukendt"
    if tool == "weather":
        city = last.get("city")
        city_text = f" (by: {city})" if city else ""
        return f"Vejrdata kom fra {source} via vejr‑værktøjet{city_text}."
    if tool == "news":
        sources = last.get("sources") or source
        return f"Nyhederne kom fra: {sources}."
    if tool == "search":
        return f"Resultaterne kom fra {source}."
    if tool == "time":
        return f"Tid/dato kommer fra {source}."
    if tool == "system":
        return f"Systeminfo kommer fra {source}."
    if tool == "ping":
        return f"Ping‑resultater kom fra {source}."
    if tool == "process":
        return f"Procesdata kom fra {source}."
    return "Jeg har ingen værktøjsdata at oplyse."


def _collect_sources(items: list[dict]) -> str:
    sources = []
    for item in items:
        src = (item.get("source") or "").strip()
        if src:
            sources.append(src)
    uniq = sorted(set(sources))
    return ", ".join(uniq) if uniq else "ukendt"


def _ticket_debug_block(tool: str, reason: str, detail: str | None, prompt: str, session_id: str | None) -> str:
    log_path = os.getenv("JARVIS_LOG_PATH", "data/logs/system.log")
    log_tail = ""
    try:
        if os.path.exists(log_path):
            with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
                lines = f.readlines()[-120:]
            log_tail = "".join(lines).strip()
    except Exception:
        log_tail = ""
    parts = [
        f"Tool: {tool}",
        f"Reason: {reason}",
        f"Detail: {detail or ''}",
        f"Prompt: {prompt}",
        f"Session: {session_id or ''}",
    ]
    if log_tail:
        parts.append("\n--- system.log (tail) ---\n" + log_tail)
    return "\n".join([p for p in parts if p])


def extract_location(prompt: str) -> str | None:
    punct = "?!.,\"()"
    stop_words = {
        "idag",
        "imorgen",
        "nu",
        "senere",
        "vejret",
        "vejr",
        "hvordan",
        "for",
        "på",
        "tak",
        "gerne",
    }
    tokens = [t.strip(punct) for t in prompt.replace(",", " ").split()]
    tokens = [t for t in tokens if t]
    if not tokens:
        return None

    for tok in tokens:
        if tok.isdigit() and len(tok) == 4:
            return tok

    start = None
    for i in range(1, len(tokens)):
        if tokens[i - 1].lower() in {"vejret", "vejr"} and tokens[i].lower() == "i":
            start = i + 1
            break

    if start is None:
        for i in range(1, len(tokens)):
            if tokens[i].upper() == "DK" and tokens[i - 1].lower() not in stop_words:
                return f"{tokens[i - 1]}, DK"
        for i in range(len(tokens) - 1):
            if tokens[i].lower() == "i":
                next_low = tokens[i + 1].lower()
                if next_low in stop_words or next_low in {"dag", "morgen"}:
                    continue
                start = i + 1
                break
        if start is None:
            return None

    loc_tokens = []
    for idx in range(start, len(tokens)):
        tok = tokens[idx]
        low = tok.lower()
        if low in stop_words:
            break
        if low == "i" and idx + 1 < len(tokens):
            next_low = tokens[idx + 1].lower()
            if next_low in {"dag", "morgen"}:
                break
        loc_tokens.append(tok)

    if not loc_tokens:
        return None
    if loc_tokens[-1].upper() == "DK" and len(loc_tokens) >= 2:
        return f"{' '.join(loc_tokens[:-1])}, DK"
    return " ".join(loc_tokens)


def run_agent(
    user_id: str,
    prompt: str,
    session_id: str | None = None,
    allowed_tools: list[str] | None = None,
    ui_city: str | None = None,
    ui_lang: str | None = None,
):
    mem = search_memory(prompt, user_id=user_id)
    session_hist = get_recent_messages(session_id, limit=8) if session_id else []
    _debug(f"🧭 run_agent: user={user_id} session={session_id} prompt={prompt!r}")
    if session_id:
        wants_prompt, custom = _session_prompt_intent(prompt)
        if wants_prompt:
            if custom.lower() in {"nulstil", "reset", "standard", "default"}:
                set_custom_prompt(session_id, None)
                return {
                    "text": "Session‑personlighed nulstillet. Jeg bruger standarden igen.",
                    "meta": {"tool_used": False},
                }
            if not custom:
                return {
                    "text": "Skriv den ønskede personlighed efter kommandoen, fx: /personlighed Kort, varm og praktisk.",
                    "meta": {"tool_used": False},
                }
            set_custom_prompt(session_id, custom)
            return {
                "text": "Session‑personlighed opdateret.",
                "meta": {"tool_used": False},
            }
    profile = get_user_profile(user_id)
    display_name = _first_name(profile, user_id)
    user_id_int = (profile or {}).get("id")
    user_key = user_id
    is_admin_user = bool((profile or {}).get("is_admin"))
    reminders_due = get_due_reminders(user_id_int) if session_id and user_id_int else []
    pending_weather = _load_state(get_pending_weather(session_id)) if session_id else {}
    pending_note = _load_state(get_pending_note(session_id)) if session_id else {}
    pending_reminder = _load_state(get_pending_reminder(session_id)) if session_id else {}
    pending_file = _load_state(get_pending_file(session_id)) if session_id else {}
    pending_image_preview = _load_state(get_pending_image_preview(session_id)) if session_id else {}
    conversation_state = (
        ConversationState.from_json(get_conversation_state(session_id)) if session_id else ConversationState()
    )
    if session_id and get_conversation_state(session_id) is None:
        set_conversation_state(session_id, conversation_state.to_json())
    mode_request = _detect_response_mode(prompt)
    if mode_request and mode_request != conversation_state.response_mode:
        conversation_state.set_response_mode(mode_request)
        if session_id:
            set_conversation_state(session_id, conversation_state.to_json())
    pending_city = None
    pending_scope = None
    pending_prompt = None
    if isinstance(pending_weather, dict) and pending_weather.get("awaiting_city"):
        pending_city = extract_location(prompt) or _simple_city(prompt)
        if pending_city:
            pending_scope = pending_weather.get("scope") or "today"
            pending_prompt = pending_weather.get("prompt") or prompt
            clear_pending_weather(session_id)
    if isinstance(pending_note, dict) and pending_note.get("awaiting_note"):
        if not _is_note_related(prompt):
            pending_note = pending_note
        else:
            clear_pending_note(session_id)
            pending_note = {}
    if isinstance(pending_reminder, dict) and pending_reminder.get("awaiting_reminder"):
        if not _is_reminder_related(prompt):
            pending_reminder = pending_reminder
        else:
            clear_pending_reminder(session_id)
            pending_reminder = {}
    file_response = handle_files(
        prompt=prompt,
        session_id=session_id,
        user_id=user_id,
        user_id_int=user_id_int,
        user_key=user_key,
        display_name=display_name,
        allowed_tools=allowed_tools,
        pending_file=pending_file,
        pending_image_preview=pending_image_preview,
        reminders_due=reminders_due,
        should_attach_reminders=_should_attach_reminders,
        prepend_reminders=_prepend_reminders,
        affirm_intent=_affirm_intent,
        deny_intent=_deny_intent,
        wants_previous_prompt=_wants_previous_prompt,
    )
    if file_response:
        return file_response

    if session_id and _resume_context_intent(prompt):
        cv_state_active = _load_state(get_cv_state(session_id))
        story_state_active = _load_state(get_story_state(session_id))
        reply = _resume_context_reply(cv_state_active, story_state_active, pending_note, pending_reminder, pending_weather)
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}

    skip_cv_intent = False
    if session_id and _cv_cancel_intent(prompt):
        set_cv_state(session_id, json.dumps({}))
        has_followup, stripped = _cancel_with_followup(prompt)
        if stripped:
            prompt = stripped
        if has_followup:
            skip_cv_intent = True
        else:
            if _cv_own_intent(prompt):
                reply = JARVIS_CV + "\nVil du have, at vi fortsætter og laver dit CV?"
            elif _cv_example_intent(prompt):
                reply = CV_STRUCTURE_TEMPLATE + "\nVil du have, at vi fortsætter og laver dit CV?"
            else:
                reply = "Forstået. Jeg lægger CV‑arbejdet på is. Hvad vil du i stedet?"
            if reminders_due and _should_attach_reminders(prompt):
                reply = _prepend_reminders(reply, reminders_due, user_id_int)
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}

    if session_id and _resume_context_intent(prompt):
        cv_state_active = _load_state(get_cv_state(session_id))
        story_state_active = _load_state(get_story_state(session_id))
        reply = _resume_context_reply(cv_state_active, story_state_active, pending_note, pending_reminder, pending_weather)
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}

    if session_id and _cv_cancel_intent(prompt):
        set_cv_state(session_id, json.dumps({}))
        has_followup, stripped = _cancel_with_followup(prompt)
        if stripped:
            prompt = stripped
        if has_followup:
            skip_cv_intent = True
        else:
            if _cv_own_intent(prompt):
                reply = JARVIS_CV + "\nVil du have, at vi fortsætter og laver dit CV?"
            elif _cv_example_intent(prompt):
                reply = CV_STRUCTURE_TEMPLATE + "\nVil du have, at vi fortsætter og laver dit CV?"
            else:
                reply = "Forstået. Jeg lægger CV‑arbejdet på is. Hvad vil du i stedet?"
            if reminders_due and _should_attach_reminders(prompt):
                reply = _prepend_reminders(reply, reminders_due, user_id_int)
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}

    if is_admin_user and prompt:
        payload, ask = _admin_create_user_from_prompt(prompt)
        if ask:
            add_memory("assistant", ask, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", ask)
            return {"text": ask, "meta": {"tool_used": False}}
        if payload:
            try:
                created = register_user(
                    payload["username"],
                    payload["password"],
                    1 if payload["is_admin"] else 0,
                    email=payload["email"],
                    full_name=payload["full_name"],
                    city=payload.get("city"),
                )
                reply = f"Bruger oprettet: {created['username']}."
            except sqlite3.IntegrityError:
                reply = "Brugernavn findes allerede. Vælg et andet."
            except Exception:
                reply = "Kunne ikke oprette brugeren lige nu."
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool_used": False}}

    mode = get_mode(session_id) if session_id else "balanced"
    # Handle history and summary intents
    from jarvis.agent_skills.history_skill import handle_history
    history_response = handle_history(user_id, prompt, session_id, allowed_tools, ui_city, ui_lang, reminders_due, user_id_int)
    if history_response:
        return history_response

    cv_state_active = _load_state(get_cv_state(session_id)) if session_id else None
    story_state_active = _load_state(get_story_state(session_id)) if session_id else None
    forced_tool = None
    resume_prompt = None
    if cv_state_active or story_state_active or pending_note or pending_reminder:
        cv_related = any(
            fn(prompt)
            for fn in [
                _cv_intent,
                _cv_example_intent,
                _cv_own_intent,
                _cv_help_intent,
                _continue_cv_intent,
                _show_cv_intent,
                _save_permanent_intent,
                _save_later_intent,
                _finalize_intent,
            ]
        )
        story_related = any(
            fn(prompt)
            for fn in [
                _story_intent,
                _continue_story_intent,
                _show_story_intent,
                _save_permanent_intent,
                _save_later_intent,
                _finalize_intent,
            ]
        )
        if cv_state_active and not cv_related:
            if _wants_weather_and_news(prompt):
                forced_tool = "weather_news"
                resume_prompt = "Skal jeg tage tråden op igen og fortsætte CV‑arbejdet, hvis De ønsker det?"
            else:
                candidate = None
                if _is_time_query(prompt):
                    candidate = "time"
                else:
                    candidate = choose_tool(prompt, allowed_tools=allowed_tools)
                    if not candidate and re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind)\b", prompt.lower()):
                        candidate = "weather"
                if candidate and (not allowed_tools or candidate in allowed_tools):
                    forced_tool = candidate
                    resume_prompt = "Skal jeg tage tråden op igen og fortsætte CV‑arbejdet, hvis De ønsker det?"
        if story_state_active and not story_related:
            if _wants_weather_and_news(prompt):
                forced_tool = "weather_news"
                resume_prompt = "Skal jeg tage tråden op igen og fortsætte historien, hvis De ønsker det?"
            else:
                candidate = None
                if _is_time_query(prompt):
                    candidate = "time"
                else:
                    candidate = choose_tool(prompt, allowed_tools=allowed_tools)
                    if not candidate and re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind)\b", prompt.lower()):
                        candidate = "weather"
                if candidate and (not allowed_tools or candidate in allowed_tools):
                    forced_tool = candidate
                    resume_prompt = "Skal jeg tage tråden op igen og fortsætte historien, hvis De ønsker det?"
        if pending_note and not _is_note_related(prompt):
            if _wants_weather_and_news(prompt):
                forced_tool = "weather_news"
                resume_prompt = "Skal jeg tage tråden op igen og fortsætte noten, hvis De ønsker det?"
            else:
                candidate = None
                if _is_time_query(prompt):
                    candidate = "time"
                else:
                    candidate = choose_tool(prompt, allowed_tools=allowed_tools)
                    if not candidate and re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind)\b", prompt.lower()):
                        candidate = "weather"
                if candidate and (not allowed_tools or candidate in allowed_tools):
                    forced_tool = candidate
                    resume_prompt = "Skal jeg tage tråden op igen og fortsætte noten, hvis De ønsker det?"
        if pending_reminder and not _is_reminder_related(prompt):
            if _wants_weather_and_news(prompt):
                forced_tool = "weather_news"
                resume_prompt = "Skal jeg tage tråden op igen og fortsætte påmindelsen, hvis De ønsker det?"
            else:
                candidate = None
                if _is_time_query(prompt):
                    candidate = "time"
                else:
                    candidate = choose_tool(prompt, allowed_tools=allowed_tools)
                    if not candidate and re.search(r"\b(vejr|vejret|vejrudsig\w*|temperatur|regn|vind)\b", prompt.lower()):
                        candidate = "weather"
                if candidate and (not allowed_tools or candidate in allowed_tools):
                    forced_tool = candidate
                    resume_prompt = "Skal jeg tage tråden op igen og fortsætte påmindelsen, hvis De ønsker det?"

    if session_id and _deny_intent(prompt):
        if cv_state_active:
            set_cv_state(session_id, json.dumps({}))
            reply = (
                "Som De ønsker. Vi parkerer CV‑arbejdet for nu. "
                "Skal vi kigge på noget andet — fx noter, vejret eller en kort opsummering?"
            )
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if session_id and _affirm_intent(prompt):
        if cv_state_active and cv_state_active.get("pending_start"):
            start_prompt = cv_state_active.get("prompt") or prompt
            cv_state = _init_cv_state(start_prompt)
            set_cv_state(session_id, json.dumps(cv_state))
            guidance = ""
            query = _extract_cv_query(start_prompt)
            search_result = tools.search_combined(query, max_items=5)
            if isinstance(search_result, dict) and search_result.get("items"):
                snippets = []
                for item in search_result["items"][:3]:
                    url = item.get("url")
                    if url:
                        article = tools.read_article(url)
                        summary = _summarize_text((article or {}).get("text") or "", sentences=1)
                        if summary:
                            snippets.append(summary)
                if snippets:
                    candidate = _summarize_text(" ".join(snippets), sentences=1)
                    if candidate and len(candidate) >= 40 and "send teksten" not in candidate.lower():
                        guidance = candidate
            first_q = _next_question(cv_state, CV_QUESTIONS)
            reply = "Fint. Jeg stiller et par korte spørgsmål."
            if guidance:
                reply += f"\nKort råd (fra søgning): {guidance}"
            if first_q:
                reply += f"\n{first_q}"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        if cv_state_active and not cv_state_active.get("done"):
            next_q = _next_question(cv_state_active, CV_QUESTIONS)
            reply = f"Naturligvis. {next_q}" if next_q else "Naturligvis. Hvad vil du tilføje?"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        if story_state_active and not story_state_active.get("done"):
            next_q = _next_question(story_state_active, STORY_QUESTIONS)
            reply = f"Naturligvis. {next_q}" if next_q else "Naturligvis. Hvad vil du justere?"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        if pending_note:
            reply = "Naturligvis. Hvad skal jeg gemme som note?"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        if pending_reminder:
            reply = "Naturligvis. Hvornår skal jeg minde dig om det?"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if session_id and _ticket_confirm_intent(prompt):
        pending = _load_state(get_ticket_state(session_id))
        if pending and user_id_int:
            ticket = _safe_create_ticket(
                user_id_int,
                pending.get("title", "Teknisk fejl"),
                pending.get("detail", ""),
                pending.get("priority", "moderat"),
            )
            set_ticket_state(session_id, json.dumps({}))
            if ticket:
                reply = f"Ticket oprettet (#{ticket['id']}) — prioritet: {ticket['priority']}."
            else:
                reply = "Jeg kunne ikke oprette ticketen lige nu."
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if session_id and _ticket_intent(prompt) and user_id_int:
        ticket = _safe_create_ticket(user_id_int, "Bruger‑ticket", prompt, "vigtig")
        reply = f"Ticket oprettet (#{ticket['id']})." if ticket else "Jeg kunne ikke oprette ticketen lige nu."
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if _name_intent(prompt):
        if profile and profile.get("full_name"):
            reply = f"Du hedder {_first_name(profile, user_id)}."
        else:
            reply = "Jeg har ikke dit navn endnu. Du kan tilføje det i Min konto."
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if session_id and _tool_source_intent(prompt):
        last = _load_state(get_last_tool(session_id))
        if not last:
            if user_id_int:
                detail = f"Manglende tool‑kontekst.\nPrompt: {prompt}"
                _safe_create_ticket(user_id_int, "Manglende tool‑kontekst", detail, "moderat")
            reply = "Jeg har ingen værktøjsdata at henvise til lige nu."
        else:
            reply = _format_tool_source(last)
            if last.get("error"):
                reply += " Sidste værktøjskald fejlede."
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": last.get("tool") if last else None, "tool_used": False}}
    if session_id and _ping_result_intent(prompt):
        last = _load_state(get_last_tool(session_id))
        if not last or last.get("tool") != "ping":
            if user_id_int:
                detail = f"Ingen ping‑resultat.\nPrompt: {prompt}"
                _safe_create_ticket(user_id_int, "Manglende ping‑resultat", detail, "moderat")
            reply = "Jeg har ingen tidligere ping‑resultater at vise."
        else:
            host = last.get("host") or "host"
            success = bool(last.get("success"))
            wants_summary_only = bool(re.search(r"\b(vis|resultat|statistik)\b", prompt.lower()))
            reply = ""
            if not wants_summary_only:
                reply = f"Ja, {host} svarer." if success else f"Nej, {host} svarer ikke."
                reply += " "
            reply += (
                f"Ping {host} {'lykkedes' if success else 'fejlede'}: "
                f"{last.get('loss_percent')}% pakketab, {last.get('avg_ms')} ms."
            )
            try:
                if float(last.get("avg_ms") or 0) >= 100:
                    reply += " Bemærk: Det er høj latency."
            except Exception:
                pass
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": "ping", "tool_used": False}}
    if session_id and _ping_count_intent(prompt):
        last = _load_state(get_last_tool(session_id))
        if not last or last.get("tool") != "ping":
            if user_id_int:
                detail = f"Ingen ping‑kontekst.\nPrompt: {prompt}"
                _safe_create_ticket(user_id_int, "Manglende ping‑kontekst", detail, "moderat")
            reply = "Jeg har ingen tidligere ping‑data at tælle."
        else:
            host = last.get("host") or "host"
            sent = last.get("sent")
            received = last.get("received")
            if sent is None or received is None:
                reply = f"Jeg sendte 5 ping til {host}."
            else:
                reply = f"Jeg sendte {sent} ping til {host}, og {received} svar kom tilbage."
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": "ping", "tool_used": False}}
    if session_id and _tool_error_intent(prompt):
        last = _load_state(get_last_tool(session_id))
        if not last or not last.get("error"):
            reply = "Jeg har ingen fejl at rapportere fra sidste værktøjskald."
        else:
            detail = last.get("detail") or "ukendt fejl"
            reply = f"Sidste fejl: {detail}"
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": last.get("tool") if last else None, "tool_used": False}}
    if session_id and _resume_context_intent(prompt):
        reply = _resume_context_reply(cv_state_active, story_state_active, pending_note, pending_reminder, pending_weather)
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if is_admin_user:
        if _admin_log_intent(prompt):
            log_path = os.getenv("JARVIS_LOG_PATH", "data/server.log")
            if not os.path.exists(log_path):
                reply = "Logfilen findes ikke endnu."
                add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": None, "tool_used": False}}
            try:
                with open(log_path, "r", encoding="utf-8", errors="ignore") as f:
                    lines = f.readlines()[-200:]
                log_text = "".join(lines).strip()
            except Exception:
                log_text = ""
            if not log_text:
                reply = "Logfilen er tom lige nu."
                add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": None, "tool_used": False}}
            if "analyser" in prompt.lower() or "find fejl" in prompt.lower():
                messages = [
                    {"role": "system", "content": "Analyser loggen og peg på fejl og næste skridt. Dansk, kort."},
                    {"role": "assistant", "content": log_text},
                    {"role": "user", "content": "Hvad er de vigtigste fejl og hvad bør jeg gøre?"},
                ]
                res = call_ollama(messages)
                reply = res.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
                if not reply:
                    reply = "Jeg kan ikke analysere loggen lige nu."
            else:
                reply = f"Seneste loglinjer:\n{log_text}"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        ticket_reply = _ticket_reply_intent(prompt)
        if ticket_reply is not None and user_id_int:
            ticket_id, message = ticket_reply
            add_ticket_message(ticket_id, user_id_int, "admin", message)
            reply = f"Svar sendt på ticket #{ticket_id}."
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        ticket_id = _ticket_analyze_intent(prompt)
        if ticket_id is not None and user_id_int:
            ticket = get_ticket_admin(ticket_id)
            if not ticket:
                reply = "Jeg kan ikke finde den ticket."
                add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": None, "tool_used": False}}
            thread = "\n".join([f"{m['role']}: {m['content']}" for m in ticket.get("messages", [])])
            messages = [
                {"role": "system", "content": "Analyser ticket og foreslå et kort svar. Dansk."},
                {"role": "assistant", "content": thread},
                {"role": "user", "content": "Giv et forslag til svar og næste skridt."},
            ]
            res = call_ollama(messages)
            reply = res.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
            if not reply:
                reply = "Jeg kan ikke analysere ticketen lige nu."
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    # Handle history and summary intents (already handled above)
    # if _summary_intent(prompt):
    # Notes and reminders
    if _cv_help_intent(prompt):
        reply = (
            "CV‑kommandoer:\n"
            "• /cv eksempel — viser CV‑struktur + kort eksempel\n"
            "• /cv jarvis — viser JARVIS‑CV (eksempel)\n"
            "\n"
            "Vil du have, at vi fortsætter og laver dit CV?"
        )
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if _cv_example_intent(prompt):
        reply = CV_STRUCTURE_TEMPLATE + "\nVil du have, at vi fortsætter og laver dit CV?"
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if _cv_own_intent(prompt):
        reply = JARVIS_CV + "\n\n" + CV_STRUCTURE_TEMPLATE + "\nVil du have, at vi fortsætter og laver dit CV?"
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    result = handle_notes(
        prompt=prompt,
        session_id=session_id,
        user_id_int=user_id_int,
        session_hist=session_hist,
        reminders_due=reminders_due,
        format_dt=_format_dt,
        format_note_brief=_format_note_brief,
        should_attach_reminders=_should_attach_reminders,
        prepend_reminders=_prepend_reminders,
    )
    if result:
        return result
    if session_id and _cv_intent(prompt) and not skip_cv_intent and not _load_state(get_cv_state(session_id)):
        set_cv_state(session_id, json.dumps({"pending_start": True, "prompt": prompt}, ensure_ascii=False))
        reply = "Vil du have, at jeg hjælper dig med et CV? Svar ja/nej."
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if session_id and _story_intent(prompt) and not _load_state(get_story_state(session_id)):
        topic = _extract_story_topic(prompt)
        wants_style = _story_needs_questions(prompt)
        fmt = _detect_format(prompt) or ""
        story_state = {
            "step": 0,
            "answers": {},
            "done": False,
            "needs_research": wants_style,
            "format": fmt if fmt else None,
            "auto_finalize": _save_text_intent(prompt),
            "persist": _save_permanent_intent(prompt),
        }
        if topic:
            story_state["answers"]["topic"] = topic
            story_state["step"] = 1
        set_story_state(session_id, json.dumps(story_state))
        if not wants_style and topic:
            story_prompt = _story_prompt_from_state(story_state)
            system = (
                "Du skriver en kort dansk historie. Brug kun info i prompten. "
                "Skriv flydende og sammenhængende."
            )
            story_messages = [
                {"role": "system", "content": system},
                {"role": "assistant", "content": story_prompt},
                {"role": "user", "content": "Skriv historien."},
            ]
            res = call_ollama(story_messages)
            story_text = res.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
            if not story_text:
                story_text = "Jeg kunne ikke skrive historien lige nu."
            story_state["draft"] = story_text
            story_state["done"] = True
            set_story_state(session_id, json.dumps(story_state))
            if story_state.get("auto_finalize") and story_state.get("format"):
                fmt = story_state.get("format") or "txt"
                temp = not story_state.get("persist")
                filename = _write_text_file(user_id, story_text, fmt, "tekst", temp=temp)
                if filename:
                    url = _make_download_link(user_id_int, session_id, filename, temp)
                    reply = f"{story_text}\n\nHer er din tekst: {_wrap_download_link(url)}\n{_download_notice()}"
                    if temp:
                        reply += "\nFilen slettes automatisk efter download, medmindre du beder mig gemme den."
                    data = {"type": "file", "title": "Tekst", "label": "Download tekst", "url": url}
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "data": data, "meta": {"tool": None, "tool_used": False}}
            reply = story_text + "\n\nVil du have den gemt som pdf, docx eller txt?"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        first_q = _next_question(story_state, STORY_QUESTIONS)
        reply = "Jeg kan skrive en historie eller en stil. Jeg stiller et par spørgsmål."
        if first_q:
            reply += f"\n{first_q}"
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if session_id:
        if _cv_cancel_intent(prompt):
            set_cv_state(session_id, json.dumps({}))
            if _cv_own_intent(prompt):
                reply = JARVIS_CV + "\nVil du have, at vi fortsætter og laver dit CV?"
            elif _cv_example_intent(prompt):
                reply = CV_STRUCTURE_TEMPLATE + "\nVil du have, at vi fortsætter og laver dit CV?"
            else:
                reply = "Forstået. Jeg lægger CV‑arbejdet på is. Hvad vil du i stedet?"
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        if prompt.lower().strip().startswith("annuller historie"):
            set_story_state(session_id, json.dumps({}))
            reply = "Historie-flow annulleret."
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        cv_state = _load_state(get_cv_state(session_id))
        if cv_state and not forced_tool:
            if cv_state.get("draft") and not cv_state.get("finalized"):
                if _save_permanent_intent(prompt):
                    cv_state["persist"] = True
                    set_cv_state(session_id, json.dumps(cv_state))
                    reply = "Forstået. Jeg gemmer CV'et permanent, når du beder om download."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                if _save_later_intent(prompt):
                    reply = "Fint. Jeg gemmer kladden til senere. Skriv 'vis cv' når du vil se den igen."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                if _show_cv_intent(prompt):
                    reply = cv_state.get("draft", "")
                    reply += "\n\nVil du have den gemt som pdf, docx eller txt?"
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                fmt = _detect_format(prompt)
                if fmt:
                    cv_state["format"] = fmt
                    set_cv_state(session_id, json.dumps(cv_state))
                    reply = f"Format sat til {fmt.upper()}. Skriv 'gem' når du vil have download-link."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                if _finalize_intent(prompt):
                    fmt = cv_state.get("format") or "txt"
                    temp = not cv_state.get("persist")
                    filename = _write_cv_file(user_id, cv_state.get("draft", ""), fmt, temp=temp)
                    if not filename:
                        reply = "Jeg kan ikke gemme i det format. Vælg pdf, docx eller txt."
                        add_message(session_id, "assistant", reply)
                        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                    url = _make_download_link(user_id_int, session_id, filename, temp)
                    cv_state["finalized"] = True
                    cv_state["file"] = filename
                    set_cv_state(session_id, json.dumps(cv_state))
                    reply = f"Her er dit CV: {_wrap_download_link(url)}\n{_download_notice()}"
                    if temp:
                        reply += "\nFilen slettes automatisk efter download, medmindre du beder mig gemme den."
                    data = {"type": "file", "title": "CV", "label": "Download CV", "url": url}
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "data": data, "meta": {"tool": None, "tool_used": False}}
                if _continue_cv_intent(prompt):
                    cv_state["edit_mode"] = True
                    set_cv_state(session_id, json.dumps(cv_state))
                    reply = "Skriv hvad du vil ændre eller tilføje."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                if cv_state.get("edit_mode"):
                    extra = cv_state.setdefault("answers", {}).get("other", "")
                    cv_state["answers"]["other"] = (extra + "\n" + prompt).strip()
                    cv_state["edit_mode"] = False
                    set_cv_state(session_id, json.dumps(cv_state))
                    reply = "Noteret. Skriv 'gem' for download-link eller vælg format (pdf/docx/txt)."
                    add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": None, "tool_used": False}}
            if cv_state.get("step", 0) >= 0 and not cv_state.get("done"):
                cv_state = _update_state(cv_state, prompt, CV_QUESTIONS)
                next_q = _next_question(cv_state, CV_QUESTIONS)
                if next_q:
                    set_cv_state(session_id, json.dumps(cv_state))
                    reply = next_q
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                cv_state["done"] = True
                profile_name = (profile or {}).get("full_name") or user_id
                cv_prompt = _cv_prompt_from_state(cv_state)
                system = (
                    "Du skriver et professionelt dansk CV. Brug kun info i prompten. "
                    "Formatér med overskrifter og punktform. Vær kort og præcis."
                )
                cv_messages = [
                    {"role": "system", "content": system},
                    {"role": "assistant", "content": f"Navn: {profile_name}\n{cv_prompt}"},
                    {"role": "user", "content": "Skriv et CV-kladdetekst."},
                ]
                res = call_ollama(cv_messages)
                cv_text = res.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
                if not cv_text:
                    cv_text = "Jeg kunne ikke generere et CV lige nu."
                cv_state["draft"] = cv_text
                cv_state["finalized"] = False
                set_cv_state(session_id, json.dumps(cv_state))
                if cv_state.get("auto_finalize") and cv_state.get("format"):
                    fmt = cv_state.get("format") or "txt"
                    temp = not cv_state.get("persist")
                    filename = _write_cv_file(user_id, cv_text, fmt, temp=temp)
                    if filename:
                        url = _make_download_link(user_id_int, session_id, filename, temp)
                        cv_state["finalized"] = True
                        cv_state["file"] = filename
                        set_cv_state(session_id, json.dumps(cv_state))
                        reply = f"{cv_text}\n\nHer er dit CV: {_wrap_download_link(url)}\n{_download_notice()}"
                        if temp:
                            reply += "\nFilen slettes automatisk efter download, medmindre du beder mig gemme den."
                        data = {"type": "file", "title": "CV", "label": "Download CV", "url": url}
                        add_message(session_id, "assistant", reply)
                        return {"text": reply, "data": data, "meta": {"tool": None, "tool_used": False}}
                reply = cv_text + "\n\nVil du have det gemt som pdf, docx eller txt?"
                add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        if _show_cv_intent(prompt):
            reply = "Jeg har ingen CV-kladde endnu. Start med at skrive, at du vil lave et CV."
            add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": None, "tool_used": False}}
        story_state = _load_state(get_story_state(session_id))
        if story_state:
            if story_state.get("draft") and not story_state.get("finalized"):
                if _save_permanent_intent(prompt):
                    story_state["persist"] = True
                    set_story_state(session_id, json.dumps(story_state))
                    reply = "Forstået. Jeg gemmer teksten permanent, når du beder om download."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                if _save_later_intent(prompt):
                    reply = "Fint. Jeg gemmer kladden til senere. Skriv 'vis tekst' for at få den igen."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                fmt = _detect_format(prompt)
                if fmt:
                    story_state["format"] = fmt
                    set_story_state(session_id, json.dumps(story_state))
                    reply = f"Format sat til {fmt.upper()}. Skriv 'gem' når du vil have download-link."
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                if _finalize_intent(prompt):
                    fmt = story_state.get("format") or "txt"
                    temp = not story_state.get("persist")
                    filename = _write_text_file(user_id, story_state.get("draft", ""), fmt, "tekst", temp=temp)
                    if not filename:
                        reply = "Jeg kan ikke gemme i det format. Vælg pdf, docx eller txt."
                        add_message(session_id, "assistant", reply)
                        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                    url = _make_download_link(user_id_int, session_id, filename, temp)
                    story_state["finalized"] = True
                    story_state["file"] = filename
                    set_story_state(session_id, json.dumps(story_state))
                    reply = f"Her er din tekst: {_wrap_download_link(url)}\n{_download_notice()}"
                    if temp:
                        reply += "\nFilen slettes automatisk efter download, medmindre du beder mig gemme den."
                    data = {"type": "file", "title": "Tekst", "label": "Download tekst", "url": url}
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "data": data, "meta": {"tool": None, "tool_used": False}}
                if prompt.lower().strip().startswith("vis tekst"):
                    reply = story_state.get("draft", "")
                    reply += "\n\nVil du have den gemt som pdf, docx eller txt?"
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
            if story_state.get("step", 0) >= 0 and not story_state.get("done"):
                story_state = _update_state(story_state, prompt, STORY_QUESTIONS)
                next_q = _next_question(story_state, STORY_QUESTIONS)
                if next_q:
                    set_story_state(session_id, json.dumps(story_state))
                    reply = next_q
                    add_message(session_id, "assistant", reply)
                    return {"text": reply, "meta": {"tool": None, "tool_used": False}}
                story_state["done"] = True
                set_story_state(session_id, json.dumps(story_state))
                story_prompt = _story_prompt_from_state(story_state)
                research = ""
                if story_state.get("needs_research"):
                    topic = story_state.get("answers", {}).get("topic")
                    if topic:
                        search_result = tools.search_combined(topic, max_items=5)
                        facts = []
                        for item in (search_result or {}).get("items", [])[:3]:
                            url = item.get("url")
                            if url:
                                article = tools.read_article(url)
                                summary = _summarize_text((article or {}).get("text") or "", sentences=1)
                                if summary:
                                    facts.append(summary)
                        if facts:
                            research = "Fakta (fra søgning): " + " ".join(facts)
                system = (
                    "Du skriver en dansk historie eller stil. Brug kun info i prompten. "
                    "Hvis der er fakta, brug kun dem. Skriv flydende og sammenhængende."
                )
                story_messages = [
                    {"role": "system", "content": system},
                    {"role": "assistant", "content": f"{story_prompt}\n{research}".strip()},
                    {"role": "user", "content": "Skriv teksten."},
                ]
                res = call_ollama(story_messages)
                story_text = res.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
                if not story_text:
                    story_text = "Jeg kunne ikke skrive teksten lige nu."
                story_state["draft"] = story_text
                story_state["finalized"] = False
                set_story_state(session_id, json.dumps(story_state))
                if story_state.get("auto_finalize") and story_state.get("format"):
                    fmt = story_state.get("format") or "txt"
                    temp = not story_state.get("persist")
                    filename = _write_text_file(user_id, story_text, fmt, "tekst", temp=temp)
                    if filename:
                        url = _make_download_link(user_id_int, session_id, filename, temp)
                        reply = f"{story_text}\n\nHer er din tekst: {_wrap_download_link(url)}\n{_download_notice()}"
                        if temp:
                            reply += "\nFilen slettes automatisk efter download, medmindre du beder mig gemme den."
                        data = {"type": "file", "title": "Tekst", "label": "Download tekst", "url": url}
                        add_message(session_id, "assistant", reply)
                        return {"text": reply, "data": data, "meta": {"tool": None, "tool_used": False}}
                reply = story_text + "\n\nVil du have den gemt som pdf, docx eller txt?"
                add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    if _farewell_intent(prompt):
        name = _first_name(profile, "")
        if name:
            reply = f"Tak for i dag, {name}. Det bliver mig en fornøjelse at hjælpe igen i morgen."
        else:
            reply = "Tak for i dag. Det bliver mig en fornøjelse at hjælpe igen i morgen."
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": None, "tool_used": False}}
    idx = _read_news_index(prompt)
    if idx is not None and not session_id:
        reply = "Læs-kommando kræver en aktiv session."
        return {"text": reply}
    if session_id:
        if idx is not None:
            _debug(f"📰 read-request: session={session_id} idx={idx}")
            raw = get_last_news(session_id)
            source = "news"
            if not raw:
                raw = get_last_search(session_id)
                source = "search"
            if not raw:
                _debug("📰 read-request: no last_news stored")
                reply = "Jeg har ingen gemte resultater i denne session endnu."
                add_message(session_id, "assistant", reply)
                return {"text": reply}
            try:
                payload = json.loads(raw)
            except Exception:
                _debug("📰 read-request: last results JSON parse failed")
                reply = "Jeg kunne ikke læse den seneste liste."
                add_message(session_id, "assistant", reply)
                return {"text": reply}
            items = payload.get("items", [])
            if idx < 1 or idx > len(items):
                reply = f"Nummer {idx} findes ikke i listen."
                add_message(session_id, "assistant", reply)
                return {"text": reply}
            url = items[idx - 1].get("url")
            cached_summary = items[idx - 1].get("summary")
            if source == "search" and cached_summary:
                reply = cached_summary
                add_message(session_id, "assistant", reply)
                return {"text": reply}
            article = tools.read_article(url)
            if article.get("error") or not article.get("text"):
                _debug(f"📰 read-request: article fetch failed: {article.get('error')}")
                reply = "Jeg kan ikke hente artiklen lige nu."
                add_message(session_id, "assistant", reply)
                return {"text": reply}
            rule = "Opsummér kun ud fra teksten. Hvis tekst mangler, sig det."
            article_messages = [
                {"role": "system", "content": f"{SYSTEM_PROMPT}\n{rule}"},
                {"role": "assistant", "content": article.get("text", "")},
                {"role": "user", "content": "Giv et kort resume."},
            ]
            res = call_ollama(article_messages)
            reply = res.get("choices",[{}])[0].get("message",{}).get("content","")
            reply = _dedupe_repeated_words(reply)
            if not reply.strip():
                reply = "Jeg kan ikke opsummere artiklen lige nu."
            add_message(session_id, "assistant", reply)
            return {
                "text": reply,
                "data": {
                    "type": "article",
                    "url": article.get("url"),
                    "title": article.get("title") or "Artikel",
                },
            }
    if session_id and prompt.strip().lower().startswith("/mode "):
        requested = prompt.strip().split(None, 1)[1].strip().lower()
        if requested in {"fakta", "snak", "balanced"}:
            set_mode(session_id, requested)
            reply = f"Mode sat til {requested}."
        else:
            reply = "Ukendt mode. Brug /mode fakta, /mode snak eller /mode balanced."
        add_message(session_id, "assistant", reply)
        return {"text": reply}
    if not session_id and prompt.strip().lower().startswith("/mode "):
        return {"text": "Mode kræver en session."}

    if allowed_tools is not None:
        requested_tool = choose_tool(prompt, allowed_tools=None)
        if requested_tool and requested_tool not in allowed_tools:
            label = _tool_label(requested_tool)
            reply = f"Værktøjet {label} er slået fra. Slå det til i Værktøjer og prøv igen."
            if reminders_due and _should_attach_reminders(prompt):
                reply = _prepend_reminders(reply, reminders_due, user_id_int)
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": requested_tool, "tool_used": False}}

    if session_id:
        pending_process = _load_state(get_process_state(session_id))
        pending_pid = pending_process.get("pid") if isinstance(pending_process, dict) else None
        if pending_pid and _process_confirm_intent(prompt):
            tool_result = tools.kill_process(int(pending_pid))
            set_process_state(session_id, "")
            reply = "Proces afsluttet." if tool_result.get("ok") else "Jeg kunne ikke afslutte processen."
            if reminders_due and _should_attach_reminders(prompt):
                reply = _prepend_reminders(reply, reminders_due, user_id_int)
            add_memory("assistant", reply, user_id=user_id)
            add_message(session_id, "assistant", reply)
            return {"text": reply, "data": tool_result, "meta": {"tool": "process", "tool_used": True}}

    tool = None
    cv_intent = False if skip_cv_intent else _cv_intent(prompt)
    wants_weather_news = _wants_weather_and_news(prompt)
    if forced_tool:
        tool = forced_tool
    elif pending_city:
        tool = "weather"
    elif _is_time_query(prompt) and (not allowed_tools or "time" in allowed_tools):
        tool = "time"
    elif cv_intent:
        tool = "search" if not allowed_tools or "search" in allowed_tools else None
    else:
        if wants_weather_news:
            tool = "weather_news"
        else:
            tool = choose_tool(prompt, allowed_tools=allowed_tools)
    tool_used = bool(tool)
    tool_result = None
    _debug(f"🧭 run_agent: tool={tool}")

    tool_summary = None
    if tool == "weather_news":
        weather_city = None
        base_prompt = pending_prompt or prompt
        weather_city = pending_city or extract_location(base_prompt)
        if not weather_city and session_id:
            weather_city = get_last_city(session_id)
        if not weather_city and profile:
            weather_city = (profile.get("city") or "").strip() or None
        if not weather_city and ui_city:
            weather_city = ui_city.strip() or None
        weather_reply = None
        weather_rendered = None
        weather_now = None
        weather_forecast = None
        if not weather_city:
            if session_id:
                set_pending_weather(
                    session_id,
                    json.dumps(
                        {"awaiting_city": True, "prompt": prompt, "scope": want_weather_scope(prompt)},
                        ensure_ascii=False,
                    ),
                )
            weather_reply = "Jeg mangler en by eller et postnummer for vejret."
        else:
            weather_now = tools.weather_now(weather_city)
            weather_forecast = tools.weather_forecast(weather_city)
            today_text = tools.format_weather_today(weather_now)
            tomorrow_text = tools.format_weather_tomorrow(weather_forecast)
            multi_text = tools.format_weather_5days(weather_forecast)
            scope = pending_scope or want_weather_scope(base_prompt)
            parts = []
            if scope == "today" and today_text:
                parts.append(("I dag", today_text))
            if scope == "tomorrow" and tomorrow_text:
                parts.append(("I morgen", tomorrow_text))
            if scope == "multi" and multi_text:
                parts.append(("5 dage", multi_text))
            if parts:
                name = (weather_now.get("name") if isinstance(weather_now, dict) else None) or weather_city
                header = f"{name} — i dag"
                if scope == "tomorrow":
                    header = f"{name} — i morgen"
                if scope == "multi":
                    header = f"{name} — 5 dage"
                rendered_lines = []
                for label, text in parts:
                    if label == "5 dage":
                        for sub in text.splitlines():
                            rendered_lines.append(f"• {sub}")
                    else:
                        rendered_lines.append(f"• {label}: {text}")
                weather_rendered = "\n".join([header] + rendered_lines[:5])
                weather_reply = weather_rendered
                if session_id:
                    set_last_city(session_id, name)
                    last_payload = {
                        "tool": "weather",
                        "source": "OpenWeather",
                        "city": name,
                        "scope": scope,
                    }
                    set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))

        query = _extract_news_query(prompt)
        category = "technology" if _is_tech_query(query) else None
        news_result = tools.news_combined(query, category=category)
        items = news_result.get("items", []) if isinstance(news_result, dict) else []
        if len(items) < 3:
            site_hint = (
                f"{query} site:reuters.com OR site:bbc.co.uk OR site:apnews.com "
                "OR site:theguardian.com OR site:nytimes.com"
            ).strip()
            fallback = tools.web_search_news(site_hint)
            if fallback:
                existing = {i.get("url") for i in items}
                for item in fallback:
                    if item.get("url") in existing:
                        continue
                    items.append(item)
                    existing.add(item.get("url"))
                    if len(items) >= 5:
                        break
        items = items[:5]
        for idx, item in enumerate(items, start=1):
            if not item.get("id"):
                item["id"] = f"n{idx}"
        news_result = {"type": "news", "query": query, "items": items}
        news_reply = None
        if not items:
            news_reply = "Jeg kan ikke hente nyheder lige nu. Prøv igen senere."
        else:
            intro = "Seneste nyheder" if not query else f"Seneste nyheder om {query}"
            if len(items) < 3:
                intro = f"{intro} (jeg fandt kun {len(items)} relevante lige nu)"
            lines = _format_news_items(items)
            news_reply = intro + ":\n" + "\n".join(lines)
            news_reply += "\n\nVil du have mig til at læse en af dem? Skriv: læs nr 2."
            if session_id:
                set_last_news(session_id, json.dumps(news_result, ensure_ascii=False))
                last_payload = {
                    "tool": "news",
                    "source": _collect_sources(items),
                    "sources": _collect_sources(items),
                    "count": len(items),
                }
                set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))

        reply_parts = []
        if weather_reply:
            reply_parts.append(weather_reply)
        if news_reply:
            reply_parts.append(news_reply)
        reply_body = "\n\n".join([p for p in reply_parts if p])
        if weather_reply and news_reply:
            reply = "Selvfølgelig — her er både vejret og nyhederne.\n\n" + reply_body
        else:
            reply = reply_body
        if not reply:
            reply = "Jeg kan ikke hente data lige nu. Prøv igen om lidt."
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        combined_data = {
            "type": "mixed",
            "weather": {
                "location": (weather_now.get("name") if isinstance(weather_now, dict) else None) or weather_city,
                "now": weather_now,
                "forecast": weather_forecast,
                "scope": pending_scope or want_weather_scope(base_prompt),
            } if weather_now or weather_forecast else None,
            "news": news_result if items else None,
        }
        return {
            "text": reply,
            "rendered_text": weather_rendered,
            "data": combined_data,
            "meta": {"tool": "weather+news", "tool_used": True},
        }

    if tool == "weather":
        base_prompt = pending_prompt or prompt
        city = pending_city or extract_location(base_prompt)
        if not city and session_id:
            city = get_last_city(session_id)
        if not city and profile:
            city = (profile.get("city") or "").strip() or None
        if not city and ui_city:
            city = ui_city.strip() or None
        if not city:
            if session_id:
                set_pending_weather(
                    session_id,
                    json.dumps(
                        {"awaiting_city": True, "prompt": prompt, "scope": want_weather_scope(prompt)},
                        ensure_ascii=False,
                    ),
                )
            reply = "Jeg mangler en by eller et postnummer. Hvis du ønsker det, kan jeg bruge din profilby."
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": "weather", "tool_used": False}}
        now = tools.weather_now(city)
        forecast = tools.weather_forecast(city)
        tool_result = {
            "now": now,
            "forecast": forecast,
        }
        today_text = tools.format_weather_today(now)
        tomorrow_text = tools.format_weather_tomorrow(forecast)
        multi_text = tools.format_weather_5days(forecast)
        scope = pending_scope or want_weather_scope(base_prompt)
        parts = []
        if scope == "today" and today_text:
            parts.append(("I dag", today_text))
        if scope == "tomorrow" and tomorrow_text:
            parts.append(("I morgen", tomorrow_text))
        if scope == "multi" and multi_text:
            parts.append(("5 dage", multi_text))
        tool_summary = parts if parts else None
    elif tool == "news":
        query = _extract_news_query(prompt)
        category = "technology" if _is_tech_query(query) else None
        tool_result = tools.news_combined(query, category=category)
    elif tool == "search":
        query = _extract_cv_query(prompt) if cv_intent else _extract_search_query(prompt)
        tool_result = tools.search_combined(query, max_items=5)
    elif tool == "currency":
        tool_result = tools.currency_convert("EUR","DKK")
    elif tool == "time":
        tool_result = tools.time_now()
    elif tool == "system":
        tool_result = tools.system_info()
    elif tool == "ping":
        host = _extract_host(prompt)
        tool_result = tools.ping_host(host) if host else {"error": "missing_host"}
    elif tool == "process":
        action = _process_action(prompt)
        if action == "kill":
            pid = _extract_pid(prompt)
            if not pid:
                tool_result = {"error": "missing_pid"}
            elif not session_id:
                reply = "Proces-afslutning kræver en aktiv session."
                add_memory("assistant", reply, user_id=user_id)
                return {"text": reply, "meta": {"tool": "process", "tool_used": False}}
            else:
                set_process_state(session_id, json.dumps({"pid": pid}))
                reply = f"Jeg kan afslutte proces {pid}. Skriv 'bekræft' for at fortsætte."
                if reminders_due and _should_attach_reminders(prompt):
                    reply = _prepend_reminders(reply, reminders_due, user_id_int)
                add_memory("assistant", reply, user_id=user_id)
                add_message(session_id, "assistant", reply)
                return {"text": reply, "meta": {"tool": "process", "tool_used": False}}
        elif action == "find":
            tool_result = tools.find_process(prompt)
        else:
            tool_result = tools.list_processes(10)

    failure = _tool_failed(tool, tool_result)
    if failure:
        reason, detail = failure
        reply = _tool_failure_reply(tool, reason, detail)
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        if session_id:
            last_payload = {
                "tool": tool,
                "source": "ukendt",
                "error": True,
                "detail": f"{reason}: {detail or ''}".strip(),
            }
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        if session_id and user_id_int:
            detail_text = _ticket_debug_block(tool, reason, detail, prompt, session_id)
            ticket = _safe_create_ticket(user_id_int, f"Fejl i {tool}", detail_text, _ticket_priority(detail_text))
            if ticket:
                reply += f"\nJeg har oprettet en ticket til dev‑teamet (#{ticket['id']})."
            else:
                reply += "\nJeg kunne ikke oprette en ticket lige nu."
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply}

    if tool == "currency" and session_id:
        last_payload = {"tool": "currency", "source": "exchangerate.host"}
        set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))

    if tool == "news":
        query = tool_result.get("query", "") if isinstance(tool_result, dict) else ""
        items = tool_result.get("items", []) if isinstance(tool_result, dict) else []
        if len(items) < 3:
            site_hint = (
                f"{query} site:reuters.com OR site:bbc.co.uk OR site:apnews.com "
                "OR site:theguardian.com OR site:nytimes.com"
            ).strip()
            fallback = tools.web_search_news(site_hint)
            if fallback:
                existing = {i.get("url") for i in items}
                for item in fallback:
                    if item.get("url") in existing:
                        continue
                    items.append(item)
                    existing.add(item.get("url"))
                    if len(items) >= 5:
                        break
            tool_result = {
                "type": "news",
                "query": query,
                "items": items[:5],
            }
        items = items[:5]
        for idx, item in enumerate(items, start=1):
            if not item.get("id"):
                item["id"] = f"n{idx}"
        tool_result = {"type": "news", "query": query, "items": items}
        if not items:
            reply = "Jeg kan ikke hente nyheder lige nu. Prøv igen senere."
            if resume_prompt:
                reply += f"\n\n{resume_prompt}"
            if reminders_due and _should_attach_reminders(prompt):
                reply = _prepend_reminders(reply, reminders_due, user_id_int)
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {
                "text": reply,
                "data": tool_result,
                "meta": {"status": "empty", "tool": "news", "tool_used": True},
            }
        intro = "Seneste nyheder" if not query else f"Seneste nyheder om {query}"
        if len(items) < 3:
            intro = f"{intro} (jeg fandt kun {len(items)} relevante lige nu)"
        lines = _format_news_items(items)
        reply = intro + ":\n" + "\n".join(lines)
        reply += "\n\nVil du have mig til at læse en af dem? Skriv: læs nr 2."
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        if session_id:
            set_last_news(session_id, json.dumps(tool_result, ensure_ascii=False))
            last_payload = {
                "tool": "news",
                "source": _collect_sources(items),
                "sources": _collect_sources(items),
                "count": len(items),
            }
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        status = "partial" if len(items) < 3 else "ok"
        return {
            "text": reply,
            "data": tool_result,
            "meta": {"status": status, "count": len(items), "tool": "news", "tool_used": True},
        }

    if tool == "search":
        if not isinstance(tool_result, dict) or tool_result.get("error"):
            reply = "Jeg kan ikke hente søgeresultater lige nu. Prøv igen senere."
            if resume_prompt:
                reply += f"\n\n{resume_prompt}"
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {"text": reply, "meta": {"tool": "search", "tool_used": True}}
        items = tool_result.get("items", [])
        if not items:
            reply = "Jeg fandt ingen resultater."
            if resume_prompt:
                reply += f"\n\n{resume_prompt}"
            if reminders_due and _should_attach_reminders(prompt):
                reply = _prepend_reminders(reply, reminders_due, user_id_int)
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {"text": reply, "data": tool_result, "meta": {"tool": "search", "tool_used": True}}
        max_items = 5
        items = items[:max_items]
        deep = _is_deep_search(prompt)
        sentence_count = 3 if deep else 1
        enriched = []
        for item in items:
            url = item.get("url")
            summary = ""
            if url:
                article = tools.read_article(url)
                text = (article or {}).get("text") or ""
                summary = _summarize_text(text, sentences=sentence_count)
                if not summary:
                    summary = _extract_excerpt(text)
            if summary:
                item = dict(item)
                item["summary"] = summary
            enriched.append(item)
        items = enriched
        if len(items) < 5:
            intro = f"Jeg fandt kun {len(items)} relevante resultater."
        else:
            intro = "Her er de mest relevante resultater:"
        summaries = [i.get("summary") for i in items if i.get("summary")]
        snippets = [i.get("snippet") for i in items if i.get("snippet")]
        description = ""
        if summaries:
            description = _summarize_text(" ".join(summaries), sentences=2 if deep else 1)
        elif snippets:
            description = _shorten(" ".join(snippets[:2]), 220)
        lines = _format_search_items(items)
        reply_parts = []
        if cv_intent:
            summary_system = (
                "Du er en CV-assistent. Svar på dansk og brug kun info i teksten. "
                "Giv en kort og praktisk vejledning, der matcher brugerens ønskede job. "
                "Ingen gæt."
            )
            summary_messages = [
                {"role": "system", "content": summary_system},
                {"role": "assistant", "content": "\n".join(summaries or snippets)},
                {"role": "user", "content": "Giv en samlet beskrivelse og en CV-struktur (overskrifter) samt 3-6 konkrete punkter."},
            ]
            res = call_ollama(summary_messages)
            cv_reply = res.get("choices", [{}])[0].get("message", {}).get("content", "")
            cv_reply = _dedupe_repeated_words(cv_reply).strip()
            if cv_reply:
                reply_parts.append(cv_reply)
        if description:
            label = "Mere dybdegående beskrivelse" if deep else "Kort beskrivelse"
            reply_parts.append(f"{label} (ud fra resultaterne): {description}")
        reply_parts.append(intro)
        reply_parts.append("\n".join(lines))
        reply = "\n".join([p for p in reply_parts if p])
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        tool_result = {"type": "search", "query": tool_result.get("query", ""), "items": items}
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        if session_id:
            set_last_search(session_id, json.dumps(tool_result, ensure_ascii=False))
            last_payload = {
                "tool": "search",
                "source": _collect_sources(items),
                "count": len(items),
            }
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply, "data": tool_result, "meta": {"tool": "search", "tool_used": True, "intent": "cv" if cv_intent else "search"}}

    if tool == "time":
        now_iso = tool_result if isinstance(tool_result, str) else None
        try:
            dt = datetime.fromisoformat(now_iso) if now_iso else datetime.now(timezone.utc)
            lang = (ui_lang or "").lower()
            if not lang:
                prompt_l = prompt.lower()
                en_markers = ["what time", "time is", "what is the time", "what is the date", "date is", "updated"]
                lang = "en" if any(marker in prompt_l for marker in en_markers) else "da"
            if lang.startswith("en"):
                if _is_date_query(prompt):
                    reply = f"Today is {dt.strftime('%d.%m.%Y')}."
                else:
                    reply = f"The time is {dt.strftime('%H:%M')}."
            else:
                if _is_date_query(prompt):
                    reply = f"I dag er det {dt.strftime('%d.%m.%Y')}."
                else:
                    reply = f"Klokken er {dt.strftime('%H:%M')}."
        except Exception:
            lang = (ui_lang or "").lower()
            reply = "I can't fetch the time right now." if lang.startswith("en") else "Jeg kan ikke hente tiden lige nu."
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        if session_id:
            last_payload = {"tool": "time", "source": "system_time"}
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": "time", "tool_used": True}}

    if tool == "system":
        if not isinstance(tool_result, dict) or tool_result.get("error"):
            reply = "Jeg kan ikke hente systeminfo lige nu."
        else:
            reply = _format_system_info(tool_result, prompt)
        reply = f"{_butler_prefix(display_name)}\n{reply}"
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        if session_id:
            last_payload = {"tool": "system", "source": "system_probe"}
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply, "data": tool_result, "meta": {"tool": "system", "tool_used": True}}

    if tool == "ping":
        if not isinstance(tool_result, dict) or tool_result.get("error"):
            reply = "Jeg kan ikke køre ping lige nu. Angiv et hostnavn eller IP."
        else:
            host = tool_result.get("host")
            success = bool(tool_result.get("success"))
            wants_details = bool(re.search(r"\b(detaljer|resultat|pinget|statistik|ms|latens)\b", prompt.lower()))
            if success:
                reply = f"Ja, {host} svarer."
            else:
                reply = f"Nej, {host} svarer ikke."
            if wants_details:
                status = "lykkedes" if success else "fejlede"
                reply += (
                    f" Ping {host} {status}: "
                    f"{tool_result.get('loss_percent')}% pakketab, {tool_result.get('avg_ms')} ms."
                )
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        if session_id:
            last_payload = {
                "tool": "ping",
                "source": "system_ping",
                "host": tool_result.get("host") if isinstance(tool_result, dict) else None,
                "success": tool_result.get("success") if isinstance(tool_result, dict) else None,
                "loss_percent": tool_result.get("loss_percent") if isinstance(tool_result, dict) else None,
                "avg_ms": tool_result.get("avg_ms") if isinstance(tool_result, dict) else None,
                "sent": tool_result.get("sent") if isinstance(tool_result, dict) else None,
                "received": tool_result.get("received") if isinstance(tool_result, dict) else None,
            }
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply, "data": tool_result, "meta": {"tool": "ping", "tool_used": True}}

    if tool == "process":
        if not isinstance(tool_result, dict) or tool_result.get("error"):
            reply = "Jeg kan ikke hente processer lige nu."
        elif tool_result.get("type") == "kill":
            reply = "Proces afsluttet." if tool_result.get("ok") else "Jeg kunne ikke afslutte processen."
        else:
            items = tool_result.get("items", [])
            if not items:
                reply = "Ingen processer fundet."
            else:
                sorted_items = sorted(items, key=lambda x: x.get("cpu", 0), reverse=True)
                match = _find_process_match(prompt, items)
                if match and _process_analysis_intent(prompt):
                    name = match.get("name") or "processen"
                    pid = match.get("pid")
                    cpu = match.get("cpu")
                    details = {
                        "uvicorn": "Uvicorn er en ASGI‑server, typisk brugt til FastAPI‑apps.",
                        "ollama": "Ollama er model‑serveren, der kører LLM‑inference lokalt.",
                        "firefox": "Firefox er en webbrowser.",
                        "code": "VS Code er en editor/IDE.",
                        "python": "Python er fortolkeren, som kører scripts og services.",
                        "node": "Node.js kører JavaScript‑services.",
                        "postgres": "PostgreSQL er en database‑server.",
                        "redis": "Redis er en hurtig nøgle‑/værdi‑database.",
                        "docker": "Docker kører containere og relaterede processer.",
                        "nginx": "Nginx er en webserver/reverse proxy.",
                        "gnome-shell": "GNOME Shell er skrivebordsmiljøets proces.",
                    }
                    desc = details.get((name or "").lower(), "Det er en systemproces/applikation.")
                    ownership = "Det ligner sandsynligvis Jarvis‑serveren her på maskinen." if (name or "").lower() == "uvicorn" else "Ud fra navnet alene kan jeg ikke bekræfte ejerskab."
                    reply = (
                        f"Som De ønsker. {name} (PID {pid}) bruger {cpu}% CPU.\n"
                        f"{desc}\n{ownership}"
                    )
                else:
                    top = sorted_items[0]
                    head = f"Jeg bemærker, at {top['name']} (PID {top['pid']}) ligger højest med {top['cpu']}% CPU."
                    lines = [
                        f"{i['pid']} {i['name']} — CPU {i['cpu']}% • MEM {i['mem']}%"
                        for i in sorted_items[:5]
                    ]
                    reply = f"{head}\n\nTop 5 processer:\n" + "\n".join(lines)
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        if session_id:
            last_payload = {"tool": "process", "source": "system_process"}
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply, "data": tool_result, "meta": {"tool": "process", "tool_used": True}}

    if tool == "weather":
        if not tool_summary:
            reply = "Jeg kan ikke hente data lige nu. Prøv igen om lidt."
            if resume_prompt:
                reply += f"\n\n{resume_prompt}"
            add_memory("assistant", reply, user_id=user_id)
            if session_id:
                add_message(session_id, "assistant", reply)
            return {"text": reply}
        if session_id:
            stored_city = now.get("name") if isinstance(now, dict) else None
            set_last_city(session_id, stored_city or city)
        name = (now.get("name") if isinstance(now, dict) else None) or city
        scope = want_weather_scope(prompt)
        header = f"{name} — i dag"
        if scope == "tomorrow":
            header = f"{name} — i morgen"
        if scope == "multi":
            header = f"{name} — 5 dage"
        lines = [header]
        rendered_lines = []
        for label, text in tool_summary:
            if label == "5 dage":
                for sub in text.splitlines():
                    rendered_lines.append(f"• {sub}")
            else:
                rendered_lines.append(f"• {label}: {text}")
        rendered_text = "\n".join([header] + rendered_lines[:5])
        reply = rendered_text
        if mode == "snak":
            reply = f"{reply}\nBritisk vejr, dansk tålmodighed."
        if session_id:
            last_payload = {
                "tool": "weather",
                "source": "OpenWeather",
                "city": name,
                "scope": scope,
            }
            set_last_tool(session_id, json.dumps(last_payload, ensure_ascii=False))
        if resume_prompt:
            reply += f"\n\n{resume_prompt}"
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        if os.getenv("TTS", "true").lower() == "true":
            audio_file = tts.speak(reply)
            if audio_file:
                return {
                    "text": reply,
                    "rendered_text": rendered_text,
                    "data": {
                        "type": "weather",
                        "location": name,
                        "now": now,
                        "forecast": forecast,
                        "scope": scope,
                    },
                    "meta": {"tool": "weather", "tool_used": True},
                    "audio": audio_file,
                }
        return {
            "text": reply,
            "rendered_text": rendered_text,
            "data": {
                "type": "weather",
                "location": name,
                "now": now,
                "forecast": forecast,
                "scope": scope,
            },
            "meta": {"tool": "weather", "tool_used": True},
        }

    mode_hint = ""
    if mode == "fakta":
        mode_hint = "Mode fakta: max 4 linjer, ingen humor, punktform."
    elif mode == "snak":
        mode_hint = "Mode snak: max 6 linjer, en kort tør joke til sidst."
    else:
        mode_hint = "Mode balanced: kort og præcist."

    sys_prompt = _get_system_prompt()
    if session_id:
        override = get_custom_prompt(session_id)
        if override:
            sys_prompt = override
    name_hint = f"Brugerens navn er {display_name}."
    messages = [{"role": "system", "content": f"{sys_prompt}\n{name_hint}\n{mode_hint}"}]
    if mem:
        messages.append({"role": "assistant", "content": "\n".join(mem)})
    messages.extend(_format_history(session_hist))
    messages.append({"role": "user", "content": prompt})
    if tool_result is not None:
        payload = tool_summary if tool_summary else tool_result
        messages.append({"role": "assistant", "content": f"Tool result: {payload}"})

    res = call_ollama(messages)
    if res.get("error"):
        reply = "Beklager — modellen svarede ikke i tide. Prøv igen, eller brug et kortere spørgsmål."
        if reminders_due and _should_attach_reminders(prompt):
            reply = _prepend_reminders(reply, reminders_due, user_id_int)
        add_memory("assistant", reply, user_id=user_id)
        if session_id:
            add_message(session_id, "assistant", reply)
        return {"text": reply, "meta": {"tool": tool or None, "tool_used": tool_used}}
    if os.getenv("DEBUG_OLLAMA") == "1":
        print(f"DEBUG_OLLAMA response keys: {list(res.keys())}")
    reply = res.get("choices",[{}])[0].get("message",{}).get("content","")
    if os.getenv("DEBUG_OLLAMA") == "1":
        print(f"DEBUG_OLLAMA reply length: {len(reply)}")
    reply = _dedupe_repeated_words(reply)
    if not reply or not reply.strip():
        print("⚠ Empty model reply; returning fallback message")
        reply = "Beklager — jeg fik et tomt svar fra modellen. Prøv igen."
    if resume_prompt:
        reply = f"{reply}\n\n{resume_prompt}"

    if reminders_due and _should_attach_reminders(prompt):
        reply = _prepend_reminders(reply, reminders_due, user_id_int)
    add_memory("assistant", reply, user_id=user_id)
    if session_id:
        add_message(session_id, "assistant", reply)
        conversation_state.update_summary(reply)
        set_conversation_state(session_id, conversation_state.to_json())

    if os.getenv("TTS", "true").lower() == "true":
        audio_file = tts.speak(reply)
        if audio_file:
            return {"text": reply, "audio": audio_file, "meta": {"tool": tool or None, "tool_used": tool_used}}
        return {"text": reply, "meta": {"tool": tool or None, "tool_used": tool_used}}

    return {"text": reply, "meta": {"tool": tool or None, "tool_used": tool_used}}
